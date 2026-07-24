from django.db import connection
from django.db.models import Count, Sum, Q
from django.utils import timezone
from rest_framework import viewsets, status
from rest_framework.decorators import api_view, permission_classes, action
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from datetime import date, datetime
import os, io, requests, openpyxl
from concurrent.futures import ThreadPoolExecutor
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt
from django.http import FileResponse, HttpResponse
from django.conf import settings

from .models import Viloyat, Tuman, Mahalla, Hisobot, Rasm, TargibotUtkazilganJoy, AuditLog, Inspektor, MurojaatUsul, MurojaatKasb, Murojaat, KunlikIshlar, HamkorTashkilot, HamkorXodim, ViloyatInfratuzilma
from .serializers import (ViloyatSerializer, TumanSerializer, MahallaSerializer,
                          HisobotSerializer, TargibotJoySerializer, InspektorSerializer,
                          MurojaatUsulSerializer, MurojaatKasbSerializer, MurojaatSerializer,
                          KunlikIshlarSerializer, HamkorTashkilotSerializer, HamkorXodimSerializer,
                          ViloyatInfratuzilmaSerializer)
from .permissions import IsRespublika, IsViloyatOrAbove
from apps.accounts.models import User
from apps.accounts.serializers import FoydalanuvchiSerializer

BOT_URL = f"https://api.telegram.org/bot{settings.TELEGRAM_TOKEN}"

# ── Telegram yordamchi ────────────────────────────────────────────────────────
def tg_send(chat_id, text, reply_to=None):
    data = {'chat_id': chat_id, 'text': text, 'parse_mode': 'HTML'}
    if reply_to:
        data['reply_to_message_id'] = reply_to
    try:
        requests.post(f"{BOT_URL}/sendMessage", data=data, timeout=5)
    except Exception:
        pass

# ── Audit log yordamchi ───────────────────────────────────────────────────────
def audit(request, amal, tavsif=''):
    try:
        AuditLog.objects.create(user=request.user, amal=amal, tavsif=tavsif)
    except Exception:
        pass

# ── Excel eksport yordamchi ───────────────────────────────────────────────────
def excel_response(headers, rows, filename):
    """Chiroyli Excel fayl qaytaradi."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Hisobot'

    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill('solid', fgColor='3730A3')  # indigo
    center      = Alignment(horizontal='center', vertical='center')

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font      = header_font
        cell.fill      = header_fill
        cell.alignment = center

    alt_fill = PatternFill('solid', fgColor='EEF2FF')
    for row_idx, row in enumerate(rows, 2):
        for col_idx, val in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            if row_idx % 2 == 0:
                cell.fill = alt_fill

    # Ustun kengliklari
    for col in ws.columns:
        max_len = max((len(str(c.value or '')) for c in col), default=0)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 40)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    resp = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    resp['Content-Disposition'] = f'attachment; filename="{filename}"'
    return resp

# ── Viloyat filter yordamchilari ──────────────────────────────────────────────
def get_viloyat_qs_filter(request, prefix='mahalla__tuman__viloyat_id'):
    """
    ORM queryset uchun viloyat filteri.
    prefix:
      Hisobot  → 'mahalla__tuman__viloyat_id' (default)
      Mahalla  → 'tuman__viloyat_id'
      Tuman    → 'viloyat_id'
    """
    if request.user.role == 'viloyat':
        return {prefix: request.user.viloyat_id}
    vid = request.GET.get('viloyat')
    return {prefix: int(vid)} if vid else {}

def get_viloyat_sql(request):
    """
    Raw SQL uchun (extra_where, extra_params).
    extra_where: ' AND tuman.viloyat_id = %s' yoki ''
    """
    if request.user.role == 'viloyat':
        return ' AND tuman.viloyat_id = %s', [request.user.viloyat_id]
    vid = request.GET.get('viloyat')
    if vid:
        return ' AND tuman.viloyat_id = %s', [int(vid)]
    return '', []

# ── Respublika Dashboard ──────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated, IsRespublika])
def respublika_dashboard(request):
    today = date.today().isoformat()
    viloyatlar = Viloyat.objects.all()
    result = []
    for v in viloyatlar:
        qs = Hisobot.objects.filter(mahalla__tuman__viloyat_id=v.id, mahalla__is_viloyat=True)
        admin = User.objects.filter(viloyat_id=v.id, role='viloyat').first()
        result.append({
            'id':                v.id,
            'nomi':              v.nomi,
            'mahalla_soni':      Mahalla.objects.filter(tuman__viloyat_id=v.id, is_tuman=False).count(),
            'yangi':             qs.filter(status=1).count(),
            'tasdiqlangan':      qs.filter(status=2).count(),
            'rad_etilgan':       qs.filter(status=3).count(),
            'bugun_tasdiqlangan':qs.filter(status=2, qushilgan_vaqt__date=today).count(),
            'admin_username':    admin.username if admin else None,
        })
    resp_qs = Hisobot.objects.filter(mahalla__is_viloyat=True)
    return Response({
        'jami_yangi':        resp_qs.filter(status=1).count(),
        'jami_tasdiqlangan': resp_qs.filter(status=2).count(),
        'jami_rad_etilgan':  resp_qs.filter(status=3).count(),
        'viloyatlar':        result,
    })

# ── Dashboard ─────────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def dashboard(request):
    today = date.today().isoformat()
    vf    = get_viloyat_qs_filter(request)
    qs    = Hisobot.objects.filter(**vf)

    # Disk holati (faqat respublika admin uchun)
    disk_info = {}
    if request.user.role == 'respublika':
        images_dir = os.path.join(settings.MEDIA_ROOT, 'images')
        try:
            total_bytes = sum(
                os.path.getsize(os.path.join(images_dir, f))
                for f in os.listdir(images_dir)
                if os.path.isfile(os.path.join(images_dir, f))
            )
            rasm_soni = Rasm.objects.count()
            disk_info = {
                'rasm_soni':   rasm_soni,
                'hajm_mb':     round(total_bytes / (1024 * 1024), 1),
                'hajm_gb':     round(total_bytes / (1024 ** 3), 2),
                'ortacha_kb':  round(total_bytes / max(rasm_soni, 1) / 1024, 1),
            }
        except Exception:
            disk_info = {}

    return Response({
        'yangi':        qs.filter(status=1).count(),
        'tasdiqlangan': qs.filter(status=2).count(),
        'rad_etilgan':  qs.filter(status=3).count(),
        'bugun': {
            'mfy_soni':       qs.filter(status=2, qushilgan_vaqt__date=today).values('mahalla').distinct().count(),
            'fuqarolar_soni': qs.filter(status=2, qushilgan_vaqt__date=today).aggregate(s=Sum('qatnashchilar_soni'))['s'] or 0,
        },
        'disk': disk_info,
    })

# ── Yangi targ'ibotlar ────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def yangi_targibotlar(request):
    vf   = get_viloyat_qs_filter(request)
    role = request.user.role

    # Offline/Online va OAV:
    #   respublika → barcha viloyatlardan kelgan hammasini ko'radi va tasdiqlaydi
    #   viloyat    → faqat o'z viloyatidagilarni ko'radi
    if role == 'respublika':
        # Faqat viloyat botga yuklagan targ'ibotlar (mahalla.is_viloyat=True)
        qs_reg = Hisobot.objects.filter(
            status=1, targibot_turi__in=[1, 2], mahalla__is_viloyat=True
        ).select_related('mahalla__tuman__viloyat').prefetch_related('rasmlar')
        qs_oav = Hisobot.objects.filter(
            status=1, targibot_turi__gte=3, mahalla__is_viloyat=True
        ).select_related('mahalla__tuman__viloyat').prefetch_related('rasmlar')
    elif role == 'viloyat':
        qs_reg = Hisobot.objects.filter(
            status=1, targibot_turi__in=[1, 2], **vf
        ).select_related('mahalla__tuman__viloyat').prefetch_related('rasmlar')
        qs_oav = Hisobot.objects.filter(
            status=1, targibot_turi__gte=3, **vf
        ).select_related('mahalla__tuman__viloyat').prefetch_related('rasmlar')
    else:
        qs_reg = Hisobot.objects.none()
        qs_oav = Hisobot.objects.none()

    reg_data = HisobotSerializer(qs_reg, many=True).data
    oav_data = HisobotSerializer(qs_oav, many=True).data
    return Response({'hisobotlar': reg_data, 'oav': oav_data})

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def tasdiqlash(request):
    tasdiqlangan  = request.data.get('tasdiqlangan_ids', [])
    rad           = request.data.get('rad_ids', [])
    rad_sabablar  = request.data.get('rad_sabablar', {})   # {str(id): "sabab matni"}
    vf            = get_viloyat_qs_filter(request)

    # Ruxsat: viloyat — o'z viloyatining offline/online + OAV (is_viloyat=False)
    # respublika — hammasini + OAV (is_viloyat=True)
    role = request.user.role
    base_qs = Hisobot.objects.filter(id__in=tasdiqlangan + rad)
    if role == 'viloyat':
        allowed = base_qs.filter(**vf, mahalla__is_viloyat=False).values_list('id', flat=True)
    elif role == 'respublika':
        allowed = base_qs.filter(mahalla__is_viloyat=True).values_list('id', flat=True)
    else:
        allowed = []
    allowed = set(allowed)

    t_ids = [i for i in tasdiqlangan if i in allowed]
    r_ids = [i for i in rad if i in allowed]
    Hisobot.objects.filter(id__in=t_ids).update(status=2)
    Hisobot.objects.filter(id__in=r_ids).update(status=3)

    for h in Hisobot.objects.filter(id__in=allowed).select_related('mahalla'):
        if h.status == 2:
            text = "✅ Targ'ibotingiz qabul qilindi!\nHamkorlik uchun rahmat 🙏"
        else:
            sabab = str(rad_sabablar.get(str(h.id), '')).strip()
            text  = "❌ Kechirasiz, targ'ibotingiz qabul qilinmadi."
            if sabab:
                text += f"\n\n📝 Sabab: {sabab}"

        # Yangi: Inspektor.tg_id orqali xabar
        for ins in Inspektor.objects.filter(mahalla=h.mahalla, tg_id__gt=0, is_active=True):
            tg_send(ins.tg_id, text, h.message_id)
        # Legacy: Mahalla.tg_id (faqat active inspektor bo'lmasa)
        if h.mahalla.tg_id > 0 and not Inspektor.objects.filter(mahalla=h.mahalla, tg_id__gt=0, is_active=True).exists():
            tg_send(h.mahalla.tg_id, text, h.message_id)

    # ── Rad etilgan hisobotlar rasmlarini diskdan o'chirish ──────────────────
    if r_ids:
        for h_rad in Hisobot.objects.filter(id__in=r_ids).prefetch_related('rasmlar'):
            for r in h_rad.rasmlar.all():
                fpath = os.path.join(settings.MEDIA_ROOT, 'images', r.rasm_url)
                try:
                    if os.path.exists(fpath):
                        os.remove(fpath)
                except Exception:
                    pass
                r.delete()

    if t_ids:
        audit(request, 'tasdiqlash', f"{len(t_ids)} ta hisobot tasdiqlandi: {t_ids}")
    if r_ids:
        sabablar_txt = ', '.join(
            f"{i}:{rad_sabablar.get(str(i),'—')}" for i in r_ids
        )
        audit(request, 'rad_etish', f"{len(r_ids)} ta rad etildi. Sabablar: {sabablar_txt}")
    return Response({'ok': True})

# ── Tasdiqlangan ──────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def tasdiqlangan(request):
    # Oy boshi → bugun default
    today_str     = date.today().isoformat()
    month_start   = date.today().replace(day=1).isoformat()
    start         = request.GET.get('start', month_start)
    end           = request.GET.get('end',   today_str)
    tuman_id      = request.GET.get('tuman')
    mahalla_id    = request.GET.get('mahalla')
    search        = request.GET.get('q', '').strip()

    vf = get_viloyat_qs_filter(request)
    qs = Hisobot.objects.filter(
        status=2,
        qushilgan_vaqt__date__gte=start,
        qushilgan_vaqt__date__lte=end,
        **vf
    ).select_related('mahalla__tuman__viloyat').prefetch_related('rasmlar')
    if request.user.role == 'respublika':
        qs = qs.filter(mahalla__is_viloyat=True)

    if tuman_id:
        qs = qs.filter(mahalla__tuman_id=tuman_id)
    if mahalla_id:
        qs = qs.filter(mahalla_id=mahalla_id)
    if search:
        qs = qs.filter(mahalla__mahalla_nomi__icontains=search)

    # Limit: ko'p ma'lumot bo'lsa sekin bo'lmasin (default 500)
    limit = min(int(request.GET.get('limit', 500) or 500), 1000)
    return Response(HisobotSerializer(qs[:limit], many=True).data)

@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def rasmlarni_ochir(request):
    start = request.GET.get('start')
    end   = request.GET.get('end')
    vf    = get_viloyat_qs_filter(request)
    qs    = Hisobot.objects.filter(status=2, qushilgan_vaqt__date__gte=start,
                                   qushilgan_vaqt__date__lte=end, **vf)
    for h in qs:
        for r in h.rasmlar.all():
            path = os.path.join(settings.MEDIA_ROOT, 'images', r.rasm_url)
            if os.path.exists(path):
                os.remove(path)
            r.delete()
    return Response({'ok': True})

# ── Rad etilgan ───────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def rad_etilgan(request):
    from datetime import timedelta
    default_start = (date.today() - timedelta(days=30)).isoformat()
    start  = request.GET.get('start', default_start)
    end    = request.GET.get('end',   date.today().isoformat())
    vf     = get_viloyat_qs_filter(request)
    qs     = Hisobot.objects.filter(status=3, qushilgan_vaqt__date__gte=start,
                                    qushilgan_vaqt__date__lte=end, **vf
                    ).select_related('mahalla__tuman__viloyat').prefetch_related('rasmlar')
    if request.user.role == 'respublika':
        qs = qs.filter(mahalla__is_viloyat=True)
    limit = min(int(request.GET.get('limit', 500) or 500), 1000)
    return Response(HisobotSerializer(qs[:limit], many=True).data)

# ── Hisobot (mahalla bo'yicha) ────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def hisobot(request):
    start = request.GET.get('start', date.today().isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    extra_where, extra_params = get_viloyat_sql(request)
    sql = f"""
        SELECT tuman.tuman_nomi, mahalla.mahalla_nomi, mahalla.inspektor_fio, mahalla.inspektor_tel,
               (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s AND h.targibot_turi=1) AS offline_targibot_soni,
               (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s AND h.targibot_turi=1) AS offline_qatnashchilar,
               (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s AND h.targibot_turi=2) AS online_targibot_soni,
               (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s AND h.targibot_turi=2) AS online_qatnashchilar,
               (SELECT COUNT(*) FROM murojaat m WHERE m.mahalla_id=mahalla.id
                AND m.sana BETWEEN %s AND %s) AS murojaat_soni
        FROM mahalla JOIN tuman ON mahalla.tuman_id=tuman.id
        WHERE 1=1{extra_where}
        ORDER BY tuman.id, mahalla.mahalla_nomi
    """
    params = [start, end] * 4 + [start, end] + extra_params
    with connection.cursor() as cur:
        cur.execute(sql, params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    for r in rows:
        r['offline_targibot_soni'] = int(r['offline_targibot_soni'] or 0)
        r['online_targibot_soni']  = int(r['online_targibot_soni']  or 0)
        r['offline_qatnashchilar'] = int(r['offline_qatnashchilar'] or 0)
        r['online_qatnashchilar']  = int(r['online_qatnashchilar']  or 0)
        r['jami_fuqarolar']        = r['offline_qatnashchilar'] + r['online_qatnashchilar']
        r['murojaat_soni']         = int(r['murojaat_soni'] or 0)

    if request.GET.get('excel'):
        headers = ['#', 'Tuman', 'Mahalla', 'Inspektor FIO', 'Telefon',
                   'Offline targ.', 'Offline fuk.', 'Online targ.', 'Online fuk.', 'Jami fuqarolar', 'Murojaat']
        data = [[i+1, r['tuman_nomi'], r['mahalla_nomi'], r['inspektor_fio'], r['inspektor_tel'],
                 r['offline_targibot_soni'], r['offline_qatnashchilar'],
                 r['online_targibot_soni'], r['online_qatnashchilar'], r['jami_fuqarolar'], r['murojaat_soni']]
                for i, r in enumerate(rows)]
        audit(request, 'excel_yuklab_olish', f"Mahallalar hisoboti {start}–{end}")
        return excel_response(headers, data, f"hisobot_mahallalar_{start}_{end}.xlsx")

    return Response(rows)

# ── Hisobot kunlik ────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def hisobot_kunlik(request):
    kun     = request.GET.get('start', request.GET.get('kun', date.today().isoformat()))
    hafta   = datetime.strptime(kun, '%Y-%m-%d').weekday()
    php_day = (hafta + 1) % 7

    extra_where, extra_params = get_viloyat_sql(request)
    vf_mahalla = get_viloyat_qs_filter(request, 'tuman__viloyat_id')

    if request.GET.get('warning'):
        mahallalar = Mahalla.objects.filter(
            Q(navbatchilik_kuni1=php_day) | Q(navbatchilik_kuni2=php_day),
            tg_id__gt=0, **vf_mahalla
        )
        targets = []
        for m in mahallalar:
            bor = Hisobot.objects.filter(mahalla=m, qushilgan_vaqt__date=kun).exclude(status=0).exists()
            if not bor:
                targets.append(m)

        def _send(m):
            tg_send(m.tg_id, "⚠️ Eslatma: Bugungi targ'ibot hisobotini yubormagansiz!")

        with ThreadPoolExecutor(max_workers=10) as ex:
            list(ex.map(_send, targets))

        if targets:
            ids = [m.id for m in targets]
            with connection.cursor() as cur:
                placeholders = ','.join(['%s'] * len(ids))
                cur.execute(
                    f"UPDATE mahalla SET ogohlantirish_kun=%s WHERE id IN ({placeholders})",
                    [kun] + ids
                )

        return Response({'ok': True, 'yuborildi': len(targets)})

    sql = f"""
        SELECT tuman.tuman_nomi, mahalla.mahalla_nomi, mahalla.inspektor_fio, mahalla.inspektor_tel,
               (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=1) AS offline_soni,
               (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=1) AS offline_qatnashchi,
               (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=2) AS online_soni,
               (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=2) AS online_qatnashchi,
               (SELECT COUNT(*) FROM hisobot h2 WHERE h2.mahalla_id=mahalla.id
                AND DATE(h2.qushilgan_vaqt)=%s AND h2.status IN (1,2,3)) AS yuborilgan_soni,
               (mahalla.ogohlantirish_kun = %s) AS ogohlantirish_yuborildi,
               (SELECT COALESCE(SUM(h.offline_18_gacha),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s) AS offline_18_gacha,
               (SELECT COALESCE(SUM(h.offline_18_katta),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s) AS offline_18_katta,
               (SELECT COALESCE(SUM(h.online_18_gacha),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s) AS online_18_gacha,
               (SELECT COALESCE(SUM(h.online_18_katta),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s) AS online_18_katta
        FROM mahalla JOIN tuman ON mahalla.tuman_id=tuman.id
        WHERE (mahalla.navbatchilik_kuni1=%s OR mahalla.navbatchilik_kuni2=%s){extra_where}
        ORDER BY tuman.id, mahalla.mahalla_nomi
    """
    with connection.cursor() as cur:
        cur.execute(sql, [kun, kun, kun, kun, kun, kun, kun, kun, kun, kun, php_day, php_day] + extra_params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    for r in rows:
        r['yuborilgan']             = bool(r.pop('yuborilgan_soni', 0))
        r['ogohlantirish_yuborildi'] = bool(r.get('ogohlantirish_yuborildi', 0))
        r['jami_fuqarolar']         = int(r['offline_qatnashchi'] or 0) + int(r['online_qatnashchi'] or 0)
    return Response(rows)

# ── Hisobot tumanlar ──────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def hisobot_tumanlar(request):
    start = request.GET.get('start', date.today().isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    extra_where, extra_params = get_viloyat_sql(request)
    sql = f"""
        SELECT tuman.tuman_nomi,
               (SELECT COUNT(*) FROM mahalla m WHERE m.tuman_id=tuman.id AND m.is_tuman=0) AS mahalla_soni,
               COALESCE(SUM(h.targibot_turi=1), 0)  AS offline_targibot_soni,
               COALESCE(SUM(CASE WHEN h.targibot_turi=1 THEN h.qatnashchilar_soni ELSE 0 END),0) AS offline_qatnashchilar,
               COALESCE(SUM(h.targibot_turi=2), 0)  AS online_targibot_soni,
               COALESCE(SUM(CASE WHEN h.targibot_turi=2 THEN h.qatnashchilar_soni ELSE 0 END),0) AS online_qatnashchilar,
               COALESCE(SUM(h.offline_18_gacha),0) AS offline_18_gacha,
               COALESCE(SUM(h.offline_18_katta),0) AS offline_18_katta,
               COALESCE(SUM(h.online_18_gacha),0)  AS online_18_gacha,
               COALESCE(SUM(h.online_18_katta),0)  AS online_18_katta,
               COALESCE(SUM(h.targibot_turi=3), 0) AS tv_soni,
               COALESCE(SUM(h.targibot_turi=4), 0) AS radio_soni,
               COALESCE(SUM(h.targibot_turi=5), 0) AS gazeta_soni,
               COALESCE(SUM(h.targibot_turi=6), 0) AS jurnal_soni,
               COALESCE(SUM(h.targibot_turi=7), 0) AS internet_soni
        FROM tuman
        LEFT JOIN mahalla ON mahalla.tuman_id=tuman.id
        LEFT JOIN hisobot h ON h.mahalla_id=mahalla.id AND h.status=2
            AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s
        WHERE 1=1{extra_where}
        GROUP BY tuman.id ORDER BY tuman.id
    """
    with connection.cursor() as cur:
        cur.execute(sql, [start, end] + extra_params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    for r in rows:
        r['offline_targibot_soni']  = int(r['offline_targibot_soni']  or 0)
        r['online_targibot_soni']   = int(r['online_targibot_soni']   or 0)
        r['offline_qatnashchilar']  = int(r['offline_qatnashchilar']  or 0)
        r['online_qatnashchilar']   = int(r['online_qatnashchilar']   or 0)
        r['jami_fuqarolar']         = r['offline_qatnashchilar'] + r['online_qatnashchilar']
        r['offline_18_gacha']       = int(r['offline_18_gacha'] or 0)
        r['offline_18_katta']       = int(r['offline_18_katta'] or 0)
        r['online_18_gacha']        = int(r['online_18_gacha']  or 0)
        r['online_18_katta']        = int(r['online_18_katta']  or 0)
        r['tv_soni']                = int(r['tv_soni']      or 0)
        r['radio_soni']             = int(r['radio_soni']   or 0)
        r['gazeta_soni']            = int(r['gazeta_soni']  or 0)
        r['jurnal_soni']            = int(r['jurnal_soni']  or 0)
        r['internet_soni']          = int(r['internet_soni'] or 0)
    return Response(rows)

# ── Hisobot viloyatlar bo'yicha (faqat respublika admin) ─────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated, IsRespublika])
def hisobot_viloyatlar(request):
    start = request.GET.get('start', date.today().isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    sql = """
        SELECT v.id, v.nomi,
               COUNT(DISTINCT t.id)  AS tuman_soni,
               (SELECT COUNT(*) FROM mahalla m2 JOIN tuman t2 ON m2.tuman_id=t2.id WHERE t2.viloyat_id=v.id AND m2.is_tuman=0) AS mahalla_soni,
               COALESCE(SUM(h.targibot_turi=1), 0) AS offline_targibot_soni,
               COALESCE(SUM(CASE WHEN h.targibot_turi=1 THEN h.qatnashchilar_soni ELSE 0 END),0) AS offline_qatnashchilar,
               COALESCE(SUM(h.targibot_turi=2), 0) AS online_targibot_soni,
               COALESCE(SUM(CASE WHEN h.targibot_turi=2 THEN h.qatnashchilar_soni ELSE 0 END),0) AS online_qatnashchilar,
               COALESCE(SUM(h.offline_18_gacha),0) AS offline_18_gacha,
               COALESCE(SUM(h.offline_18_katta),0) AS offline_18_katta,
               COALESCE(SUM(h.online_18_gacha),0)  AS online_18_gacha,
               COALESCE(SUM(h.online_18_katta),0)  AS online_18_katta,
               COALESCE(SUM(h.targibot_turi=3), 0) AS tv_soni,
               COALESCE(SUM(h.targibot_turi=4), 0) AS radio_soni,
               COALESCE(SUM(h.targibot_turi=5), 0) AS gazeta_soni,
               COALESCE(SUM(h.targibot_turi=6), 0) AS jurnal_soni,
               COALESCE(SUM(h.targibot_turi=7), 0) AS internet_soni
        FROM viloyat v
        LEFT JOIN tuman t    ON t.viloyat_id = v.id
        LEFT JOIN mahalla m  ON m.tuman_id   = t.id
        LEFT JOIN hisobot h  ON h.mahalla_id = m.id AND h.status=2
             AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s
        GROUP BY v.id, v.nomi
        ORDER BY v.id
    """
    with connection.cursor() as cur:
        cur.execute(sql, [start, end])
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]
    for r in rows:
        r['offline_targibot_soni'] = int(r['offline_targibot_soni'] or 0)
        r['online_targibot_soni']  = int(r['online_targibot_soni']  or 0)
        r['offline_qatnashchilar'] = int(r['offline_qatnashchilar'] or 0)
        r['online_qatnashchilar']  = int(r['online_qatnashchilar']  or 0)
        r['jami_fuqarolar']        = r['offline_qatnashchilar'] + r['online_qatnashchilar']
        r['offline_18_gacha']      = int(r['offline_18_gacha'] or 0)
        r['offline_18_katta']      = int(r['offline_18_katta'] or 0)
        r['online_18_gacha']       = int(r['online_18_gacha']  or 0)
        r['online_18_katta']       = int(r['online_18_katta']  or 0)
        r['tv_soni']               = int(r['tv_soni']      or 0)
        r['radio_soni']            = int(r['radio_soni']   or 0)
        r['gazeta_soni']           = int(r['gazeta_soni']  or 0)
        r['jurnal_soni']           = int(r['jurnal_soni']  or 0)
        r['internet_soni']         = int(r['internet_soni'] or 0)

    if request.GET.get('excel'):
        headers = ['#', 'Viloyat', 'Tumanlar', 'Mahallalar', 'Offline targ.', 'Offline fuk.',
                   'Online targ.', 'Online fuk.', 'Jami fuqarolar']
        data = [[i+1, r['nomi'], r['tuman_soni'], r['mahalla_soni'],
                 r['offline_targibot_soni'], r['offline_qatnashchilar'],
                 r['online_targibot_soni'], r['online_qatnashchilar'], r['jami_fuqarolar']]
                for i, r in enumerate(rows)]
        audit(request, 'excel_yuklab_olish', f"Viloyatlar hisoboti {start}–{end}")
        return excel_response(headers, data, f"hisobot_viloyatlar_{start}_{end}.xlsx")

    return Response(rows)

# ── Bugun targibot qilmaganlar ────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def qilmaganlar(request):
    today    = date.today()
    php_day  = (today.weekday() + 1) % 7
    vf       = get_viloyat_qs_filter(request, 'tuman__viloyat_id')
    mahallalar = Mahalla.objects.filter(
        Q(navbatchilik_kuni1=php_day) | Q(navbatchilik_kuni2=php_day), **vf
    ).select_related('tuman')
    result = []
    for m in mahallalar:
        bor = Hisobot.objects.filter(mahalla=m, qushilgan_vaqt__date=today).exclude(status=0).exists()
        if not bor:
            result.append({
                'id': m.id, 'mahalla_nomi': m.mahalla_nomi,
                'tuman_nomi':    m.tuman.tuman_nomi,
                'inspektor_fio': m.inspektor_fio,
                'inspektor_tel': m.inspektor_tel,
            })
    return Response(result)

# ── Xabar yuborish ────────────────────────────────────────────────────────────
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def xabar_yuborish(request):
    text = (request.data.get('matn') or request.data.get('xabar') or '').strip()
    if not text:
        return Response({'error': "Xabar bo'sh"}, status=400)
    vf_ins  = get_viloyat_qs_filter(request, 'mahalla__tuman__viloyat_id')
    vf_mal  = get_viloyat_qs_filter(request, 'tuman__viloyat_id')

    # Barcha tg_id larni yig'amiz (takrorlansiz)
    tg_ids = set()
    for ins in Inspektor.objects.filter(tg_id__gt=0, is_active=True, **vf_ins).values_list('tg_id', flat=True):
        tg_ids.add(ins)
    for tg in Mahalla.objects.filter(tg_id__gt=0, **vf_mal).values_list('tg_id', flat=True):
        tg_ids.add(tg)

    yuborildi = xato = 0
    results = {'y': 0, 'x': 0}

    def _send(tg_id):
        try:
            tg_send(tg_id, text)
            results['y'] += 1
        except Exception:
            results['x'] += 1

    with ThreadPoolExecutor(max_workers=20) as ex:
        list(ex.map(_send, tg_ids))

    yuborildi = results['y']
    xato      = results['x']
    audit(request, 'xabar_yuborish', f"{yuborildi} ta inspektorga xabar yuborildi")
    return Response({'ok': True, 'yuborildi': yuborildi, 'xato': xato})

# ── Rasmlarni siqish (admin buyrug'i) ────────────────────────────────────────
@api_view(['POST'])
@permission_classes([IsAuthenticated, IsRespublika])
def compress_images_view(request):
    """Barcha saqlangan rasmlarni siqadi. Faqat respublika admin."""
    try:
        from PIL import Image as PILImage
    except ImportError:
        return Response({'error': 'Pillow o\'rnatilmagan'}, status=500)

    quality  = int(request.data.get('quality', 70))
    max_size = int(request.data.get('max_size', 900))
    images_dir = os.path.join(settings.MEDIA_ROOT, 'images')

    processed = skipped = errors = 0
    saved_bytes = 0

    for r in Rasm.objects.all().iterator():
        fpath = os.path.join(images_dir, r.rasm_url)
        if not os.path.exists(fpath):
            skipped += 1
            continue
        try:
            size_before = os.path.getsize(fpath)
            img = PILImage.open(fpath).convert('RGB')
            img.thumbnail((max_size, max_size), PILImage.LANCZOS)
            img.save(fpath, 'JPEG', quality=quality, optimize=True)
            size_after = os.path.getsize(fpath)
            saved_bytes += max(0, size_before - size_after)
            processed += 1
        except Exception:
            errors += 1

    audit(request, 'compress_images',
          f"{processed} ta rasm siqildi, {saved_bytes//(1024*1024)} MB tejaldi")
    return Response({
        'processed':  processed,
        'skipped':    skipped,
        'errors':     errors,
        'saved_mb':   round(saved_bytes / (1024 * 1024), 1),
    })

# ── Qamrov ───────────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def qamrov(request):
    """Viloyat → Tuman → Mahalla qamrov darajasi (tasdiqlangan hisobotlar bo'yicha)."""
    today      = date.today()
    # Default: oxirgi 30 kun (joriy oy emas, chunki oy boshida ma'lumot kam bo'ladi)
    from datetime import timedelta
    default_start = (today - timedelta(days=30)).isoformat()
    start = request.GET.get('start', default_start)
    end   = request.GET.get('end',   today.isoformat())
    extra_where, extra_params = get_viloyat_sql(request)

    sql = f"""
        SELECT
            viloyat.id        AS viloyat_id,
            viloyat.nomi      AS viloyat_nomi,
            tuman.id          AS tuman_id,
            tuman.tuman_nomi  AS tuman_nomi,
            mahalla.id        AS mahalla_id,
            mahalla.mahalla_nomi,
            COUNT(h.id)       AS targibot_soni
        FROM viloyat
        JOIN tuman   ON tuman.viloyat_id  = viloyat.id
        JOIN mahalla ON mahalla.tuman_id  = tuman.id AND mahalla.is_tuman = 0
        LEFT JOIN hisobot h
               ON h.mahalla_id = mahalla.id
              AND h.status = 2
              AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s
        WHERE 1=1{extra_where}
        GROUP BY viloyat.id, viloyat.nomi,
                 tuman.id,   tuman.tuman_nomi,
                 mahalla.id, mahalla.mahalla_nomi
        ORDER BY viloyat.id, tuman.id, mahalla.mahalla_nomi
    """
    with connection.cursor() as cur:
        cur.execute(sql, [start, end] + extra_params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    # Ierarxiya qurish
    vil_dict = {}
    for r in rows:
        vid, tid = r['viloyat_id'], r['tuman_id']
        if vid not in vil_dict:
            vil_dict[vid] = {'id': vid, 'nomi': r['viloyat_nomi'], 'tum': {}}
        if tid not in vil_dict[vid]['tum']:
            vil_dict[vid]['tum'][tid] = {'id': tid, 'nomi': r['tuman_nomi'], 'mal': []}
        soni = int(r['targibot_soni'] or 0)
        vil_dict[vid]['tum'][tid]['mal'].append({
            'id': r['mahalla_id'], 'nomi': r['mahalla_nomi'],
            'targibot_soni': soni, 'qamrangan': soni > 0,
        })

    result = []
    for v in vil_dict.values():
        tumanlar, v_jami, v_qam = [], 0, 0
        for t in v['tum'].values():
            mahallalar = t['mal']
            t_jami = len(mahallalar)
            t_qam  = sum(1 for m in mahallalar if m['qamrangan'])
            v_jami += t_jami;  v_qam += t_qam
            tumanlar.append({
                'id': t['id'], 'nomi': t['nomi'],
                'jami': t_jami, 'qamrangan': t_qam, 'qolgan': t_jami - t_qam,
                'foiz': round(t_qam * 100 / t_jami) if t_jami else 0,
                'mahallalar': mahallalar,
            })
        result.append({
            'id': v['id'], 'nomi': v['nomi'],
            'jami': v_jami, 'qamrangan': v_qam, 'qolgan': v_jami - v_qam,
            'foiz': round(v_qam * 100 / v_jami) if v_jami else 0,
            'tumanlar': tumanlar,
        })
    return Response(result)

# ── GPS nuqtalar (xarita uchun) ───────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def qamrov_nuqtalar(request):
    """GPS koordinatali tasdiqlangan hisobotlar — xaritadagi pinlar uchun."""
    today      = date.today()
    from datetime import timedelta
    default_start = (today - timedelta(days=30)).isoformat()
    start = request.GET.get('start', default_start)
    end   = request.GET.get('end',   today.isoformat())
    vf    = get_viloyat_qs_filter(request)

    qs = Hisobot.objects.filter(
        status=2,
        latitude__isnull=False,
        longitude__isnull=False,
        qushilgan_vaqt__date__gte=start,
        qushilgan_vaqt__date__lte=end,
        **vf
    ).select_related('mahalla__tuman__viloyat')

    data = [{
        'id':               h.id,
        'lat':              h.latitude,
        'lng':              h.longitude,
        'mahalla_nomi':     h.mahalla.mahalla_nomi,
        'tuman_nomi':       h.mahalla.tuman.tuman_nomi,
        'targibot_turi':    h.targibot_turi,
        'qatnashchilar':    h.qatnashchilar_soni,
        'sana':             str(h.qushilgan_vaqt)[:10],
    } for h in qs]
    return Response(data)

# ── Audit Log ─────────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def audit_log_list(request):
    # Default: oxirgi 30 kun (bugun emas, chunki loglar o'tgan oyda bo'lishi mumkin)
    from datetime import timedelta
    default_end   = date.today()
    default_start = default_end - timedelta(days=30)
    start = request.GET.get('start', default_start.isoformat())
    end   = request.GET.get('end',   default_end.isoformat())
    amal  = request.GET.get('amal', '')

    qs = AuditLog.objects.filter(vaqt__date__gte=start, vaqt__date__lte=end
         ).select_related('user').order_by('-vaqt')

    if amal:
        qs = qs.filter(amal=amal)

    # Viloyat admin faqat o'z loglarini ko'radi
    if request.user.role == 'viloyat':
        qs = qs.filter(user__viloyat_id=request.user.viloyat_id)

    data = [{
        'id':       l.id,
        'user':     l.user.username if l.user else '—',
        'fish':     l.user.fish     if l.user else '—',
        'role':     l.user.role     if l.user else '—',
        'amal':     l.amal,
        'tavsif':   l.tavsif,
        'vaqt':     l.vaqt.strftime('%Y-%m-%d %H:%M:%S'),
    } for l in qs[:500]]
    return Response(data)

# ── Mahalla CRUD ──────────────────────────────────────────────────────────────
class MahallaViewSet(viewsets.ModelViewSet):
    serializer_class    = MahallaSerializer
    permission_classes  = [IsAuthenticated]

    def get_queryset(self):
        vf = get_viloyat_qs_filter(self.request, 'tuman__viloyat_id')
        return Mahalla.objects.select_related('tuman').filter(**vf)

    def perform_create(self, serializer):
        instance = serializer.save()
        audit(self.request, 'mahalla_yaratish', f"Mahalla yaratildi: {instance.mahalla_nomi}")

    def perform_update(self, serializer):
        instance = serializer.save()
        audit(self.request, 'mahalla_tahrirlash', f"Mahalla tahrirlandi: {instance.mahalla_nomi}")

    def perform_destroy(self, instance):
        audit(self.request, 'mahalla_ochirish', f"Mahalla o'chirildi: {instance.mahalla_nomi}")
        instance.delete()

# ── Tuman CRUD ────────────────────────────────────────────────────────────────
class TumanViewSet(viewsets.ModelViewSet):
    serializer_class    = TumanSerializer
    permission_classes  = [IsAuthenticated]

    def get_queryset(self):
        vf = get_viloyat_qs_filter(self.request, 'viloyat_id')
        return Tuman.objects.filter(**vf)

    def perform_create(self, serializer):
        # Viloyat admin yangi tuman yaratganda o'z viloyatini avtomatik biriktiradi
        if (self.request.user.role == 'viloyat'
                and not serializer.validated_data.get('viloyat')):
            serializer.save(viloyat_id=self.request.user.viloyat_id)
        else:
            serializer.save()

# ── Inspektor CRUD ───────────────────────────────────────────────────────────
class InspektorViewSet(viewsets.ModelViewSet):
    serializer_class   = InspektorSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = Inspektor.objects.select_related('mahalla__tuman__viloyat')
        # mahalla_id filteri
        mahalla_id = self.request.GET.get('mahalla')
        if mahalla_id:
            qs = qs.filter(mahalla_id=mahalla_id)
        # viloyat filteri
        vf = get_viloyat_qs_filter(self.request, 'mahalla__tuman__viloyat_id')
        return qs.filter(**vf)

# ── Viloyat CRUD (faqat respublika) ──────────────────────────────────────────
class ViloyatViewSet(viewsets.ModelViewSet):
    queryset            = Viloyat.objects.all()
    serializer_class    = ViloyatSerializer
    permission_classes  = [IsAuthenticated, IsRespublika]

# ── Foydalanuvchilar CRUD (faqat respublika) ──────────────────────────────────
class FoydalanuvchiViewSet(viewsets.ModelViewSet):
    queryset            = User.objects.all().order_by('id')
    serializer_class    = FoydalanuvchiSerializer
    permission_classes  = [IsAuthenticated, IsRespublika]

    def perform_create(self, serializer):
        user = serializer.save()
        parol = self.request.data.get('parol')
        if parol:
            user.set_password(parol)
            user.save()

    def perform_update(self, serializer):
        user = serializer.save()
        parol = self.request.data.get('parol')
        if parol:
            user.set_password(parol)
            user.save()

# ── Hamkor tashkilotlar CRUD ─────────────────────────────────────────────────
class HamkorTashkilotViewSet(viewsets.ModelViewSet):
    serializer_class   = HamkorTashkilotSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        vf = get_viloyat_qs_filter(self.request, 'viloyat_id')
        return HamkorTashkilot.objects.select_related('viloyat').prefetch_related('xodimlar').filter(**vf)

    def perform_create(self, serializer):
        if self.request.user.role == 'viloyat' and not serializer.validated_data.get('viloyat'):
            serializer.save(viloyat_id=self.request.user.viloyat_id)
        else:
            serializer.save()


class HamkorXodimViewSet(viewsets.ModelViewSet):
    serializer_class   = HamkorXodimSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        qs = HamkorXodim.objects.select_related('tashkilot__viloyat')
        tashkilot_id = self.request.GET.get('tashkilot')
        if tashkilot_id:
            qs = qs.filter(tashkilot_id=tashkilot_id)
        vf = get_viloyat_qs_filter(self.request, 'tashkilot__viloyat_id')
        return qs.filter(**vf)


# ── Kunlik ma'lumotnoma (arxiv) ───────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def arxiv_list(request):
    folder = os.path.join(settings.MEDIA_ROOT, 'hisobot_arxiv')
    os.makedirs(folder, exist_ok=True)
    files = sorted([f for f in os.listdir(folder) if os.path.isfile(os.path.join(folder, f))], reverse=True)
    result = []
    for f in files:
        mtime = os.path.getmtime(os.path.join(folder, f))
        result.append({
            'id':             f,
            'nomi':           f,
            'url':            f'/media/hisobot_arxiv/{f}',
            'yaratilgan_vaqt': datetime.fromtimestamp(mtime).strftime('%Y-%m-%d %H:%M'),
        })
    return Response(result)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def arxiv_yaratish(request):
    today   = date.today()
    php_day = (today.weekday() + 1) % 7
    extra_where, extra_params = get_viloyat_sql(request)
    sql = f"""
        SELECT tuman.tuman_nomi, mahalla.mahalla_nomi, mahalla.inspektor_fio, mahalla.inspektor_tel,
               (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=1) AS offline_soni,
               (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=1) AS offline_qatnashchi,
               (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=2) AS online_soni,
               (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                AND DATE(h.qushilgan_vaqt)=%s AND h.targibot_turi=2) AS online_qatnashchi
        FROM mahalla JOIN tuman ON mahalla.tuman_id=tuman.id
        WHERE (mahalla.navbatchilik_kuni1=%s OR mahalla.navbatchilik_kuni2=%s){extra_where}
    """
    with connection.cursor() as cur:
        cur.execute(sql, [today] * 4 + [php_day, php_day] + extra_params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Kunlik hisobot'
    headers = ['#','Tuman','Mahalla','Inspektor FIO','Telefon',
               'Offline soni','Offline qatnashchi','Online soni','Online qatnashchi']
    ws.append(headers)
    for i, r in enumerate(rows, 1):
        ws.append([i, r['tuman_nomi'], r['mahalla_nomi'], r['inspektor_fio'], r['inspektor_tel'],
                   r['offline_soni'], r['offline_qatnashchi'], r['online_soni'], r['online_qatnashchi']])

    folder   = os.path.join(settings.MEDIA_ROOT, 'hisobot_arxiv')
    os.makedirs(folder, exist_ok=True)
    filename = f'hisobot_{today}_{int(datetime.now().timestamp())}.xlsx'
    wb.save(os.path.join(folder, filename))
    return Response({'ok': True, 'filename': filename})

@api_view(['DELETE'])
@permission_classes([IsAuthenticated, IsRespublika])
def arxiv_ochir(request, id):
    import re
    filename = os.path.basename(id)
    # Faqat kutilgan format: hisobot_YYYY-MM-DD_timestamp.xlsx
    if not re.match(r'^hisobot_\d{4}-\d{2}-\d{2}_\d+\.xlsx$', filename):
        return Response({'error': 'Noto\'g\'ri fayl nomi'}, status=400)
    path = os.path.join(settings.MEDIA_ROOT, 'hisobot_arxiv', filename)
    if os.path.exists(path):
        audit(request, 'arxiv_ochirish', f"Fayl: {filename}")
        os.remove(path)
    return Response({'ok': True})

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def arxiv_download(request, id):
    import re
    filename = os.path.basename(id)
    if not re.match(r'^hisobot_\d{4}-\d{2}-\d{2}_\d+\.xlsx$', filename):
        return Response({'error': 'Noto\'g\'ri fayl nomi'}, status=400)
    path = os.path.join(settings.MEDIA_ROOT, 'hisobot_arxiv', filename)
    if not os.path.exists(path):
        return Response({'error': 'Fayl topilmadi'}, status=404)
    return FileResponse(open(path, 'rb'), as_attachment=True, filename=filename)

# ── Word hisobot ──────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def word_hisobot(request):
    start = request.GET.get('start', date.today().isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    vf    = get_viloyat_qs_filter(request)
    qs = Hisobot.objects.filter(status=2, qushilgan_vaqt__date__gte=start,
                                qushilgan_vaqt__date__lte=end, **vf
                 ).select_related('mahalla__tuman__viloyat')

    doc = Document()
    doc.add_heading(f"Tasdiqlangan targ'ibotlar: {start} — {end}", 0)
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['#','Mahalla','Inspektor','Turi','Qatnashchilar','Joy','Sana']):
        hdr[i].text = h

    for i, h in enumerate(qs, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = h.mahalla.mahalla_nomi
        row[2].text = h.mahalla.inspektor_fio
        row[3].text = 'Offline' if h.targibot_turi == 1 else 'Online'
        row[4].text = str(h.qatnashchilar_soni or 0)
        row[5].text = str(h.targibot_utgan_joy or '')
        row[6].text = str(h.qushilgan_vaqt)[:10]

    folder   = os.path.join(settings.MEDIA_ROOT, 'hisobot')
    os.makedirs(folder, exist_ok=True)
    filename = f'hisobot_{start}_{end}.docx'
    doc.save(os.path.join(folder, filename))
    return FileResponse(open(os.path.join(folder, filename), 'rb'),
                        as_attachment=True, filename=filename)


# ── Xavfsiz va Sog'lom Yurt – rasmiy hisobot (JSON + Excel) ─────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def hisobot_excel_xavfsiz_yurt(request):
    """
    'Xavfsiz va Sog\'lom Yurt' Excel hisobotini generatsiya qiladi.
    1-jadval: Uchrashuvlar (joylashuv kategoriyasi bo'yicha) + yoshlar
    2-jadval: OAV + Tarqatilgan materiallar + Suhbatlar
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io as _io
    from datetime import timedelta

    sana_str = request.GET.get('sana', date.today().isoformat())
    sana = datetime.strptime(sana_str, '%Y-%m-%d').date()
    oy_boshi = sana.replace(day=1).isoformat()

    sql = """
        SELECT
            v.nomi                                           AS viloyat,

            /* ── 1-jadval: Uchrashuvlar soni (kategoriya bo'yicha) ── */
            -- JAMI
            COALESCE(SUM(CASE WHEN h.targibot_turi IN (1,2)
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0)  AS j_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi IN (1,2)
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS j_oy,

            -- Mahallalar (kat=1)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=1
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS mfy_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=1
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS mfy_oy,

            -- Oliy ta'lim (kat=2)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=2
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS oliy_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=2
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS oliy_oy,

            -- Akademik litsey (kat=3)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=3
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS litsey_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=3
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS litsey_oy,

            -- O'rta ta'lim (kat=4)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=4
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS orta_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=4
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS orta_oy,

            -- Maktabgacha (kat=5)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=5
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS maktabgacha_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=5
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS maktabgacha_oy,

            -- Kasalxona (kat=6)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=6
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS kasalxona_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=6
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS kasalxona_oy,

            -- Bozorlar (kat=7)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=7
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS bozor_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=7
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS bozor_oy,

            -- HMQO (kat=8)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=8
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS hmqo_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=8
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS hmqo_oy,

            -- Boshqa (kat=9)
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=9
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS boshqa_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=1 AND j.kategoriya=9
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS boshqa_oy,

            /* ── Fuqarolar soni (offline+online jami) ── */
            COALESCE(SUM(CASE WHEN h.targibot_turi IN (1,2)
                AND DATE(h.qushilgan_vaqt)=%s THEN h.qatnashchilar_soni END),0) AS fuk_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi IN (1,2)
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.qatnashchilar_soni END),0) AS fuk_oy,

            /* ── Yoshlar bo'yicha (offline) ── */
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.offline_18_gacha END),0) AS off18g_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.offline_18_gacha END),0) AS off18g_oy,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.offline_18_katta END),0) AS off18k_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.offline_18_katta END),0) AS off18k_oy,

            /* ── Yoshlar bo'yicha (online) ── */
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.online_18_gacha END),0) AS onl18g_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.online_18_gacha END),0) AS onl18g_oy,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.online_18_katta END),0) AS onl18k_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.online_18_katta END),0) AS onl18k_oy,

            /* ── 2-jadval: OAV ── */
            COALESCE(SUM(CASE WHEN h.targibot_turi IN (3,4,5,6,7)
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS oav_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi IN (3,4,5,6,7)
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS oav_oy,

            COALESCE(SUM(CASE WHEN h.targibot_turi=3
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS tv_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=3
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS tv_oy,

            COALESCE(SUM(CASE WHEN h.targibot_turi=4
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS radio_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=4
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS radio_oy,

            COALESCE(SUM(CASE WHEN h.targibot_turi=5
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS gazeta_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=5
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS gazeta_oy,

            COALESCE(SUM(CASE WHEN h.targibot_turi=6
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS jurnal_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=6
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS jurnal_oy,

            COALESCE(SUM(CASE WHEN h.targibot_turi=7
                AND DATE(h.qushilgan_vaqt)=%s THEN 1 END),0) AS internet_bir_kun,
            COALESCE(SUM(CASE WHEN h.targibot_turi=7
                AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN 1 END),0) AS internet_oy,

            /* ── Tarqatilgan materiallar ── */
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.video_kontent_soni END),0) AS video_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.video_kontent_soni END),0) AS video_oy,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.banner_soni END),0) AS banner_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.banner_soni END),0) AS banner_oy,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.flayer_soni END),0) AS flayer_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.flayer_soni END),0) AS flayer_oy,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.buklet_soni END),0) AS buklet_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.buklet_soni END),0) AS buklet_oy,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.boshqa_material_soni END),0) AS mat_boshqa_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.boshqa_material_soni END),0) AS mat_boshqa_oy,

            /* ── Suhbatlar ── */
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt)=%s THEN h.suhbat_soni END),0) AS suhbat_bir_kun,
            COALESCE(SUM(CASE WHEN DATE(h.qushilgan_vaqt) BETWEEN %s AND %s THEN h.suhbat_soni END),0) AS suhbat_oy

        FROM viloyat v
        LEFT JOIN tuman t ON t.viloyat_id=v.id
        LEFT JOIN mahalla m ON m.tuman_id=t.id
        LEFT JOIN hisobot h ON h.mahalla_id=m.id AND h.status=2
        LEFT JOIN targibot_utkazilgan_joy j ON j.id=h.targibot_utgan_joy
        GROUP BY v.id, v.nomi
        ORDER BY v.id
    """

    # Har bir juft (bir_kun, oy) uchun parametrlar
    s = sana_str
    ob = oy_boshi
    p = []
    # j_bir_kun, j_oy
    p += [s, ob, s]
    # 9 kategoriya × (bir_kun + oy)
    for _ in range(9):
        p += [s, ob, s]
    # fuk
    p += [s, ob, s]
    # yoshlar × 4
    for _ in range(4):
        p += [s, ob, s]
    # OAV × 6 (jami+5)
    for _ in range(6):
        p += [s, ob, s]
    # Materiallar × 5
    for _ in range(5):
        p += [s, ob, s]
    # Suhbat
    p += [s, ob, s]

    with connection.cursor() as cur:
        cur.execute(sql, p)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    # Barcha qiymatlari int ga aylantirish
    for r in rows:
        for k, v in r.items():
            if k != 'viloyat':
                r[k] = int(v or 0)

    # JSON rejimi (sahifada ko'rsatish uchun)
    if request.GET.get('format') == 'json':
        return Response({'sana': sana_str, 'oy_boshi': oy_boshi, 'rows': rows})

    # Excel yaratish
    wb = Workbook()
    ws = wb.active
    ws.title = sana_str

    thin   = Side(style='thin',   color='AAAAAA')
    medium = Side(style='medium', color='1F4E79')
    brd    = Border(left=thin, right=thin, top=thin, bottom=thin)
    ctr    = Alignment(horizontal='center', vertical='center', wrap_text=True)
    lft    = Alignment(horizontal='left',   vertical='center', wrap_text=True)

    DARK   = PatternFill('solid', fgColor='1F4E79')
    MID    = PatternFill('solid', fgColor='2E75B6')
    GREEN  = PatternFill('solid', fgColor='375623')
    LT_GRN = PatternFill('solid', fgColor='E2EFDA')
    LT_BLU = PatternFill('solid', fgColor='DEEAF1')
    YELLOW = PatternFill('solid', fgColor='FFF2CC')
    ORANGE = PatternFill('solid', fgColor='FCE4D6')
    PURPLE = PatternFill('solid', fgColor='E2D9F3')

    def fnt(bold=False, color='000000', size=8, italic=False):
        return Font(name='Times New Roman', bold=bold, color=color, size=size, italic=italic)

    def hdr(ws, row, col, val, fill, bold=True, color='FFFFFF', colspan=1):
        c = ws.cell(row, col, val)
        c.font = fnt(bold, color, 8)
        c.fill = fill
        c.alignment = ctr
        c.border = brd
        if colspan > 1:
            ws.merge_cells(start_row=row, start_column=col, end_row=row, end_column=col+colspan-1)
        return c

    # ── Sarlavha ──
    ws.merge_cells('A1:BK1')
    c = ws.cell(1, 1, f"Xavfsiz va Sog'lom Yurt — {sana_str} ({oy_boshi} dan {sana_str} gacha)")
    c.font = fnt(bold=True, color='FFFFFF', size=10)
    c.fill = DARK; c.alignment = ctr
    ws.row_dimensions[1].height = 28

    # ── 1-jadval sarlavhalari (qator 2-4) ──
    # Qator 2: katta bo'limlar
    hdr(ws, 2, 1, 'Hududlar', DARK, colspan=1)
    hdr(ws, 2, 2, "O'tkazilgan uchrashuvlar soni", MID, colspan=20)
    hdr(ws, 2, 22, "Qatnashchilar soni", GREEN, colspan=10)
    hdr(ws, 2, 32, "OAV chiqishlari soni", MID, colspan=12)
    hdr(ws, 2, 44, "Tarqatilgan materiallar", ORANGE, colspan=10)
    hdr(ws, 2, 54, "Suhbatlar", PURPLE, colspan=2)

    # Qator 3: kategoriyalar
    col = 2
    for nom in ['JAMI', "Ta'lim", 'Kasalxona', 'Bozorlar',
                'Istirohat', 'Transport', 'Masjid', 'HMQO', 'Boshqa', 'Gavjum']:
        hdr(ws, 3, col, nom, LT_BLU if col == 2 else LT_GRN, color='000000', colspan=2)
        col += 2

    col = 22
    for nom in ['JAMI fuqarolar', '18 yoshgacha (off.)', '18 katta (off.)',
                '18 yoshgacha (onl.)', '18 katta (onl.)']:
        hdr(ws, 3, col, nom, LT_GRN, color='000000', colspan=2)
        col += 2

    col = 32
    for nom in ['JAMI OAV', 'Televidenie', 'Radio', 'Gazeta', 'Jurnal', 'Internet']:
        hdr(ws, 3, col, nom, LT_BLU, color='000000', colspan=2)
        col += 2

    col = 44
    for nom in ['JAMI mat.', 'Video/kontent', 'Banner', 'Flayer', 'Buklet', 'Boshqa']:
        hdr(ws, 3, col, nom, YELLOW, color='000000', colspan=2)
        col += 2

    hdr(ws, 3, 54, 'Suhbatlar', PURPLE, colspan=2)

    # Qator 4: bir kunda / oy boshidan
    ws.cell(2, 1).font = fnt(bold=True, color='FFFFFF')
    ws.merge_cells(start_row=2, start_column=1, end_row=4, end_column=1)
    for col in range(2, 56):
        c_bk = ws.cell(4, col, 'bk' if col % 2 == 0 else 'oy')
        c_bk.font = fnt(italic=True, size=7)
        c_bk.fill = LT_BLU if col < 32 else (YELLOW if col >= 44 else LT_BLU)
        c_bk.alignment = ctr; c_bk.border = brd

    ws.row_dimensions[2].height = 32
    ws.row_dimensions[3].height = 32
    ws.row_dimensions[4].height = 14

    # ── Ustun kengliklari ──
    ws.column_dimensions['A'].width = 18
    for col in range(2, 56):
        ws.column_dimensions[get_column_letter(col)].width = 6

    # ── Ma'lumot qatorlari ──
    jami = {k: 0 for k in cols[1:]}
    for ri, r in enumerate(rows, 5):
        ws.cell(ri, 1, r['viloyat']).font = fnt(bold=True)
        ws.cell(ri, 1).border = brd
        col = 2
        keys_order = [
            ('j_bir_kun','j_oy'),
            ('mfy_bir_kun','mfy_oy'),('oliy_bir_kun','oliy_oy'),
            ('litsey_bir_kun','litsey_oy'),('orta_bir_kun','orta_oy'),
            ('maktabgacha_bir_kun','maktabgacha_oy'),('kasalxona_bir_kun','kasalxona_oy'),
            ('bozor_bir_kun','bozor_oy'),('hmqo_bir_kun','hmqo_oy'),
            ('boshqa_bir_kun','boshqa_oy'),
            ('fuk_bir_kun','fuk_oy'),
            ('off18g_bir_kun','off18g_oy'),('off18k_bir_kun','off18k_oy'),
            ('onl18g_bir_kun','onl18g_oy'),('onl18k_bir_kun','onl18k_oy'),
            ('oav_bir_kun','oav_oy'),
            ('tv_bir_kun','tv_oy'),('radio_bir_kun','radio_oy'),
            ('gazeta_bir_kun','gazeta_oy'),('jurnal_bir_kun','jurnal_oy'),
            ('internet_bir_kun','internet_oy'),
            ('video_bir_kun','video_oy'),('banner_bir_kun','banner_oy'),
            ('flayer_bir_kun','flayer_oy'),('buklet_bir_kun','buklet_oy'),
            ('mat_boshqa_bir_kun','mat_boshqa_oy'),
            ('suhbat_bir_kun','suhbat_oy'),
        ]
        for k1, k2 in keys_order:
            v1 = int(r.get(k1) or 0)
            v2 = int(r.get(k2) or 0)
            jami[k1] = jami.get(k1, 0) + v1
            jami[k2] = jami.get(k2, 0) + v2
            c1 = ws.cell(ri, col, v1)
            c2 = ws.cell(ri, col+1, v2)
            for cc in (c1, c2):
                cc.font = fnt(); cc.border = brd; cc.alignment = ctr
                cc.fill = LT_BLU if col < 32 else (YELLOW if col >= 44 else LT_BLU)
            col += 2

        ws.row_dimensions[ri].height = 15

    # JAMI qatori
    ri = len(rows) + 5
    c = ws.cell(ri, 1, 'ЖАМИ')
    c.font = fnt(bold=True, color='FFFFFF'); c.fill = DARK; c.border = brd
    col = 2
    for k1, k2 in keys_order:
        for k, offset in ((k1, 0), (k2, 1)):
            cell = ws.cell(ri, col+offset, jami.get(k, 0))
            cell.font = fnt(bold=True, color='FFFFFF')
            cell.fill = DARK; cell.border = brd; cell.alignment = ctr
        col += 2
    ws.row_dimensions[ri].height = 16
    ws.freeze_panes = 'B5'

    buf = _io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"xavfsiz_soghlom_yurt_{sana_str}.xlsx"
    return HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'}
    )

# ── Xavfsiz Yurt template Excel ──────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def xavfsiz_yurt_template_excel(request):  # noqa: C901
    """
    GET /api/xavfsiz-yurt-template/?sana=YYYY-MM-DD
    Template xlsx (07.07 sheet) ga ma'lumot to'ldirib, barcha format/rang/merge saqlab qaytaradi.
    """
    import io as _io, os, copy
    from openpyxl import load_workbook
    from openpyxl.writer.excel import save_workbook
    from django.db.models import Sum as DSum

    sana_str = request.GET.get('sana', date.today().isoformat())
    sana_d   = date.fromisoformat(sana_str)
    oy_boshi = sana_d.replace(day=1).isoformat()

    # ── Oy nomi (Кирилл) ──────────────────────────────────────────────────────
    OY_NOMI = {
        1:'январь',2:'февраль',3:'март',4:'апрель',5:'май',6:'июнь',
        7:'июль',8:'август',9:'сентябрь',10:'октябрь',11:'ноябрь',12:'декабрь'
    }
    oy_nom = OY_NOMI.get(sana_d.month, str(sana_d.month))

    TEMPLATE = os.path.normpath(os.path.join(
        os.path.dirname(os.path.abspath(__file__)), '..', '..', 'xavfsiz_yurt_template.xlsx'))
    if not os.path.exists(TEMPLATE):
        return HttpResponse('Template topilmadi', status=404)

    # ── Viloyat ID → template row (07.07 sheet rows 8-21) ────────────────────
    VILOYAT_ROWS = {
        140: 8,   # Қорақалпоғистон Республикаси
        151: 9,   # Тошкент шаҳар
        154: 10,  # Тошкент вилояти
        123: 11,  # Андижон вилояти
        160: 12,  # Бухоро вилояти
        129: 13,  # Жиззах вилояти
        137: 14,  # Қашқадарё вилояти
        135: 15,  # Навоий вилояти
        132: 16,  # Наманган вилояти
        143: 17,  # Самарқанд вилояти
        145: 18,  # Сирдарё вилояти
        148: 19,  # Сурхондарё вилояти
        126: 20,  # Фарғона вилояти
        157: 21,  # Хоразм вилояти
    }
    DATA_ROWS = list(range(8, 22))
    JAMI_ROW  = 22

    # ── SQL: bir kunda (bot) ──────────────────────────────────────────────────
    sql_day = """
        SELECT v.id,
            COALESCE(SUM(CASE WHEN j.kategoriya=1  THEN 1 END),0) AS k1,
            COALESCE(SUM(CASE WHEN j.kategoriya=2  THEN 1 END),0) AS k2,
            COALESCE(SUM(CASE WHEN j.kategoriya=3  THEN 1 END),0) AS k3,
            COALESCE(SUM(CASE WHEN j.kategoriya=4  THEN 1 END),0) AS k4,
            COALESCE(SUM(CASE WHEN j.kategoriya=5  THEN 1 END),0) AS k5,
            COALESCE(SUM(CASE WHEN j.kategoriya=6  THEN 1 END),0) AS k6,
            COALESCE(SUM(CASE WHEN j.kategoriya=7  THEN 1 END),0) AS k7,
            COALESCE(SUM(CASE WHEN j.kategoriya=8  THEN 1 END),0) AS k8,
            COALESCE(SUM(CASE WHEN j.kategoriya=9  THEN 1 END),0) AS k9,
            COALESCE(SUM(CASE WHEN j.kategoriya=10 THEN 1 END),0) AS k10,
            COALESCE(SUM(CASE WHEN j.kategoriya=11 THEN 1 END),0) AS k11,
            COALESCE(SUM(CASE WHEN j.kategoriya=12 THEN 1 END),0) AS k12,
            COALESCE(SUM(CASE WHEN j.kategoriya=13 THEN 1 END),0) AS k13,
            COALESCE(SUM(h.qatnashchilar_soni),0)  AS fuk,
            COALESCE(SUM(h.offline_18_gacha),0)    AS o18g,
            COALESCE(SUM(h.offline_18_katta),0)    AS o18k,
            COALESCE(SUM(h.online_18_gacha),0)     AS n18g,
            COALESCE(SUM(h.online_18_katta),0)     AS n18k
        FROM viloyat v
        LEFT JOIN tuman t   ON t.viloyat_id=v.id
        LEFT JOIN mahalla m ON m.tuman_id=t.id
        LEFT JOIN hisobot h ON h.mahalla_id=m.id AND h.status=2
                           AND DATE(h.qushilgan_vaqt)=%s
        LEFT JOIN targibot_utkazilgan_joy j ON j.id=h.targibot_utgan_joy
        GROUP BY v.id
    """
    # ── SQL: oy boshidan (bot) ────────────────────────────────────────────────
    sql_month = """
        SELECT v.id,
            COALESCE(SUM(CASE WHEN j.kategoriya BETWEEN 1 AND 12 THEN 1 END),0) AS jami,
            COALESCE(SUM(h.qatnashchilar_soni),0) AS fuk
        FROM viloyat v
        LEFT JOIN tuman t   ON t.viloyat_id=v.id
        LEFT JOIN mahalla m ON m.tuman_id=t.id
        LEFT JOIN hisobot h ON h.mahalla_id=m.id AND h.status=2
                           AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s
        LEFT JOIN targibot_utkazilgan_joy j ON j.id=h.targibot_utgan_joy
        GROUP BY v.id
    """

    with connection.cursor() as cur:
        cur.execute(sql_day, [sana_str])
        dc = [c[0] for c in cur.description]
        bot_d = {r[0]: dict(zip(dc[1:], r[1:])) for r in cur.fetchall()}

        cur.execute(sql_month, [oy_boshi, sana_str])
        bot_m = {r[0]: {'jami': r[1], 'fuk': r[2]} for r in cur.fetchall()}

    # ── KunlikIshlar ──────────────────────────────────────────────────────────
    ki_d = {r['viloyat_id']: r for r in KunlikIshlar.objects.filter(sana=sana_str).values(
        'viloyat_id','oav_tv_soni','oav_radio_soni','oav_gazeta_jurnal_soni',
        'oav_internet_soni','oav_video_soni',
        'mat_ijtimoiy_tarmoq','mat_oz_tashabbusi','mat_flayer_buklet',
        'mat_led_ekran','mat_boshqa',
        'iio_xizmat_soni','hamkor_tashkilot_soni','sayber_soni','iio_tv_murojaati',
    )}
    ki_m = {r['viloyat_id']: r for r in KunlikIshlar.objects
        .filter(sana__gte=oy_boshi, sana__lte=sana_str).values('viloyat_id')
        .annotate(
            oav=DSum('oav_tv_soni')+DSum('oav_radio_soni')+DSum('oav_gazeta_jurnal_soni')+DSum('oav_internet_soni')+DSum('oav_video_soni'),
            mat=DSum('mat_ijtimoiy_tarmoq')+DSum('mat_oz_tashabbusi')+DSum('mat_flayer_buklet')+DSum('mat_led_ekran')+DSum('mat_boshqa'),
            iio_tv=DSum('iio_tv_murojaati'),
        )}

    infra = {i.viloyat_id: i for i in ViloyatInfratuzilma.objects.all()}

    with connection.cursor() as cur:
        cur.execute("SELECT v.id, COUNT(m.id) FROM viloyat v LEFT JOIN tuman t ON t.viloyat_id=v.id LEFT JOIN mahalla m ON m.tuman_id=t.id GROUP BY v.id")
        mfy_cnt = {r[0]: r[1] for r in cur.fetchall()}

    def g(d, k):
        return int(d.get(k) or 0) if d else 0

    # ── Template yuklash ──────────────────────────────────────────────────────
    wb = load_workbook(_io.BytesIO(open(TEMPLATE, 'rb').read()))
    SH = '07.07' if '07.07' in wb.sheetnames else wb.sheetnames[-1]
    ws = wb[SH]

    # Sarlavha oy nomini yangilash
    v1 = ws.cell(4, 3).value or ''
    ws.cell(4, 3).value = v1[:v1.lower().find('сони')+4].rstrip() + f' {oy_nom} ойи учун ' if 'сони' in v1.lower() else v1

    # ── Har viloyat uchun ma'lumot yozish ────────────────────────────────────
    for vil in Viloyat.objects.all():
        row = VILOYAT_ROWS.get(vil.id)
        if not row:
            continue

        bd  = bot_d.get(vil.id, {})
        bm  = bot_m.get(vil.id, {})
        kd  = ki_d.get(vil.id, {})
        km  = ki_m.get(vil.id, {})
        inf = infra.get(vil.id)

        def gi(k):   return g(bd, k)
        def gm(k):   return g(bm, k)
        def gk(k):   return g(kd, k)
        def gkm(k):  return g(km, k)
        def ginf(k): return int(getattr(inf, k, 0) or 0) if inf else 0

        talim_soni = ginf('oliy_talim') + ginf('akademik_litsey') + ginf('orta_talim') + ginf('maktabgacha')
        talim_birk = gi('k2') + gi('k3') + gi('k4') + gi('k5')
        jami_birk  = sum(gi(f'k{i}') for i in range(1,13)) + gk('iio_tv_murojaati')
        jami_oy    = gm('jami') + gkm('iio_tv')

        def w(col, val):
            ws.cell(row, col).value = int(val)

        # Col 2: viloyat nomi, col 34: takror nomi
        # (template da allaqachon yozilgan, o'zgartirmaymiz)

        # Аҳоли сони (col 3-5) — bizda yo'q, qoldirамiz

        # Маҳаллалар: сони=6, бир кунда=7
        w(6, mfy_cnt.get(vil.id, 0))
        w(7, gi('k1'))

        # Қизил МФЙлар: сони=8, бир кунда=9
        w(8, ginf('qizil_mfy'))
        w(9, gi('k1'))   # mahalla tashrifi = kat1

        # Таълим: сони=10, бир кунда=11
        w(10, talim_soni)
        w(11, talim_birk)

        # Касалхона: сони=12, бир кунда=13
        w(12, ginf('kasalxona'))
        w(13, gi('k6'))

        # Бозорлар: сони=14, бир кунда=15
        w(14, ginf('bozor'))
        w(15, gi('k7'))

        # Истироҳат: сони=16, бир кунда=17
        w(16, ginf('istirohat'))
        w(17, gi('k10'))

        # Жамоат транспорти: сони=18, бир кунда=19
        w(18, ginf('jamoat_transport'))
        w(19, gi('k11'))

        # Масжидлар: сони=20, бир кунда=21
        w(20, ginf('masjid'))
        w(21, gi('k12'))

        # ҲМҚО: сони=22, бир кунда=23
        w(22, ginf('xmko'))
        w(23, gi('k8'))

        # Бошқа (Telegram): сони=24, бир кунда=25
        w(24, ginf('telegram'))
        w(25, gi('k9'))

        # ИИО ТВ: ой бошидан=26, бир кунда=27
        w(26, gkm('iio_tv'))
        w(27, gk('iio_tv_murojaati'))

        # ЖАМИ учрашувлар: ой бошидан=28, бир кунда=29
        w(28, jami_oy)
        w(29, jami_birk)

        # Қамраб олинган аҳоли: jami=29(shared), қолган cols 30-33
        w(30, gi('fuk'))
        w(31, gi('o18g'))
        w(32, gi('o18k'))
        w(33, gi('n18g') + gi('n18k'))

        # OAV: жами ой бошидан=35, жами бир кунда=36
        oav_bk = gk('oav_tv_soni')+gk('oav_radio_soni')+gk('oav_gazeta_jurnal_soni')+gk('oav_internet_soni')+gk('oav_video_soni')
        w(35, gkm('oav'))
        w(36, oav_bk)
        w(37, gk('oav_tv_soni'))
        w(38, gk('oav_radio_soni'))
        w(39, gk('oav_gazeta_jurnal_soni'))
        # col 40 = Журнал — bizda gazeta+jurnal birlashgan, 0 qo'yamiz
        w(40, 0)
        w(41, gk('oav_internet_soni'))
        w(42, gk('oav_video_soni'))
        w(43, 0)  # qo'shimcha OAV turi yo'q

        # Тарқатилган материаллар: ой бошидан=44, бир кунда=45
        mat_bk = gk('mat_ijtimoiy_tarmoq')+gk('mat_oz_tashabbusi')+gk('mat_flayer_buklet')+gk('mat_led_ekran')+gk('mat_boshqa')
        w(44, gkm('mat'))
        w(45, mat_bk)
        w(46, gk('mat_ijtimoiy_tarmoq'))
        w(47, gk('mat_oz_tashabbusi'))
        w(48, gk('mat_flayer_buklet'))
        w(49, gk('mat_led_ekran'))
        w(50, gk('mat_boshqa'))

        # Пробация: сони=51, бир кунда=52
        w(51, gi('k13'))
        w(52, gi('k13'))

        # ИИО хизматлар: сони=53, бир кунда=54
        w(53, gk('iio_xizmat_soni'))
        w(54, gk('iio_xizmat_soni'))

        # Ҳамкор: col=55
        w(55, gk('hamkor_tashkilot_soni'))

        # Сайбер: col=56
        w(56, gk('sayber_soni'))

    # ── ЖАМИ qatori — template da SUM formulas bor, qoldirамiz ───────────────
    # (row 22 da allaqachon =+SUM(X8:X21) formulas mavjud)

    buf = _io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f"xavfsiz_yurt_{sana_str}.xlsx"
    return HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'}
    )


# ── Kunlik ishlar Excel export ────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated, IsViloyatOrAbove])
def kunlik_ishlar_excel(request):  # noqa: C901
    """
    GET /api/kunlik-ishlar/excel/?start=YYYY-MM-DD&end=YYYY-MM-DD
    Har kategoriya uchun: сони (infratuzilma) | бир кунда (bot)
    Ranglar: asl jadval uslubida (C6D9F0, B2A1C6, FFFFFFCC, 92D050)
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from django.db.models import Sum as DSum
    import io as _io

    start_str = request.GET.get('start', date.today().isoformat())
    end_str   = request.GET.get('end',   date.today().isoformat())
    period    = f"{start_str} — {end_str}" if start_str != end_str else start_str

    role = request.user.role
    viloyat_filter = ''
    sql_params = [start_str, end_str]
    if role == 'viloyat' and request.user.viloyat_id:
        viloyat_filter = f'AND v.id = {int(request.user.viloyat_id)}'

    # ── Bot SQL ───────────────────────────────────────────────────────────────
    sql = """
        SELECT v.id,
            COALESCE(SUM(CASE WHEN j.kategoriya=1  THEN 1 END),0) AS k1,
            COALESCE(SUM(CASE WHEN j.kategoriya=2  THEN 1 END),0) AS k2,
            COALESCE(SUM(CASE WHEN j.kategoriya=3  THEN 1 END),0) AS k3,
            COALESCE(SUM(CASE WHEN j.kategoriya=4  THEN 1 END),0) AS k4,
            COALESCE(SUM(CASE WHEN j.kategoriya=5  THEN 1 END),0) AS k5,
            COALESCE(SUM(CASE WHEN j.kategoriya=6  THEN 1 END),0) AS k6,
            COALESCE(SUM(CASE WHEN j.kategoriya=7  THEN 1 END),0) AS k7,
            COALESCE(SUM(CASE WHEN j.kategoriya=8  THEN 1 END),0) AS k8,
            COALESCE(SUM(CASE WHEN j.kategoriya=9  THEN 1 END),0) AS k9,
            COALESCE(SUM(CASE WHEN j.kategoriya=10 THEN 1 END),0) AS k10,
            COALESCE(SUM(CASE WHEN j.kategoriya=11 THEN 1 END),0) AS k11,
            COALESCE(SUM(CASE WHEN j.kategoriya=12 THEN 1 END),0) AS k12,
            COALESCE(SUM(CASE WHEN j.kategoriya=13 THEN 1 END),0) AS k13,
            COALESCE(SUM(h.qatnashchilar_soni),0) AS fuk,
            COALESCE(SUM(h.offline_18_gacha),0)   AS o18g,
            COALESCE(SUM(h.offline_18_katta),0)   AS o18k,
            COALESCE(SUM(h.online_18_gacha),0)    AS n18g,
            COALESCE(SUM(h.online_18_katta),0)    AS n18k
        FROM viloyat v
        LEFT JOIN tuman t   ON t.viloyat_id=v.id
        LEFT JOIN mahalla m ON m.tuman_id=t.id
        LEFT JOIN hisobot h ON h.mahalla_id=m.id AND h.status=2
                           AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s
        LEFT JOIN targibot_utkazilgan_joy j ON j.id=h.targibot_utgan_joy
        WHERE 1=1 """ + viloyat_filter + """
        GROUP BY v.id ORDER BY v.id
    """
    with connection.cursor() as cur:
        cur.execute(sql, [start_str, end_str])
        dc = [c[0] for c in cur.description]
        bot = {r[0]: dict(zip(dc[1:], r[1:])) for r in cur.fetchall()}

    ki_agg = (KunlikIshlar.objects
        .filter(sana__gte=start_str, sana__lte=end_str, status=3).values('viloyat_id')
        .annotate(
            tv=DSum('oav_tv_soni'), radio=DSum('oav_radio_soni'),
            gazeta=DSum('oav_gazeta_jurnal_soni'), internet=DSum('oav_internet_soni'),
            video=DSum('oav_video_soni'), v10k=DSum('oav_video_10k'),
            v100k=DSum('oav_video_100k'), v1m=DSum('oav_video_1m'),
            ijt=DSum('mat_ijtimoiy_tarmoq'), ozt=DSum('mat_oz_tashabbusi'),
            flay=DSum('mat_flayer_buklet'), led=DSum('mat_led_ekran'),
            matb=DSum('mat_boshqa'), iiox=DSum('iio_xizmat_soni'),
            hmkr=DSum('hamkor_tashkilot_soni'), sayb=DSum('sayber_soni'),
            iiotv=DSum('iio_tv_murojaati'),
        ))
    ki = {r['viloyat_id']: {k: int(v or 0) for k,v in r.items() if k!='viloyat_id'} for r in ki_agg}

    infra = {i.viloyat_id: i for i in ViloyatInfratuzilma.objects.all()}
    with connection.cursor() as cur:
        cur.execute("SELECT v.id, COUNT(m.id) FROM viloyat v LEFT JOIN tuman t ON t.viloyat_id=v.id LEFT JOIN mahalla m ON m.tuman_id=t.id GROUP BY v.id")
        mfy_cnt = {r[0]: r[1] for r in cur.fetchall()}

    viloyatlar = list(Viloyat.objects.all().order_by('nomi'))

    # ── Ranglar (asl jadval uslubi) ───────────────────────────────────────────
    wb = Workbook()
    ws = wb.active
    ws.title = period[:31]

    thin = Side(style='thin', color='AAAAAA')
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)
    ctr  = Alignment(horizontal='center', vertical='center', wrap_text=True)
    lft  = Alignment(horizontal='left',   vertical='center', wrap_text=True)

    # Asl jadval ranglari
    HDR1   = PatternFill('solid', fgColor='0070C0')   # Ko'k — asosiy sarlavha
    HDR2   = PatternFill('solid', fgColor='C6D9F0')   # Och ko'k — kategoriya sarlavha
    SONI   = PatternFill('solid', fgColor='B2A1C6')   # Och binafsha — "сони" ustun
    BIRK   = PatternFill('solid', fgColor='FFFFFF')   # Oq — "бир кунда" ustun
    YELLOW = PatternFill('solid', fgColor='FFFFCC')   # Sariq — ЖАМИ
    GREEN  = PatternFill('solid', fgColor='92D050')   # Yashil — IIO TV
    QATN   = PatternFill('solid', fgColor='E2EFDA')   # Och yashil — qatnashchilar
    OAV_H  = PatternFill('solid', fgColor='C6D9F0')   # OAV sarlavha
    OAV_D  = PatternFill('solid', fgColor='FFFFFF')
    MAT_H  = PatternFill('solid', fgColor='FCE4D6')   # Materiallar sarlavha
    MAT_D  = PatternFill('solid', fgColor='FFFFFF')
    PROB   = PatternFill('solid', fgColor='E2D9F3')   # Probatsiya
    JAMI_F = PatternFill('solid', fgColor='0070C0')   # ЖАМИ qatori
    GRAY   = PatternFill('solid', fgColor='F2F2F2')   # Hudud ustun

    def fnt(bold=False, color='000000', size=9):
        return Font(name='Times New Roman', bold=bold, color=color, size=size)

    def h(row, col, val, fill, bold=True, fc='000000', span=1, sz=9, rows=1):
        c = ws.cell(row, col, val)
        c.font = fnt(bold, fc, sz)
        c.fill = fill; c.alignment = ctr; c.border = brd
        if span > 1 or rows > 1:
            ws.merge_cells(start_row=row, start_column=col,
                           end_row=row+rows-1, end_column=col+span-1)
        return c

    # ═══════════════════════════════════════════════════════════════════════════
    # USTUN TUZILISHI:
    # 1:№  2:Ҳудуд
    # Uchrashuvlar — har kategoriya 2 ustun: сони | бир кунда
    #   3-4:  Маҳаллалар
    #   5-6:  Қизил МФЙлар
    #   7-8:  Таълим муассасалари
    #   9-10: Касалхона
    #  11-12: Бозорлар
    #  13-14: Истироҳат боғлари
    #  15-16: Жамоат транспорти
    #  17-18: Масжидлар
    #  19-20: ҲМҚО
    #  21-22: Бошқа (Telegram)
    #  23-24: ИИО ТВ мурожаати (сони=0, бир кунда)
    #  25: ЖАМИ учрашувлар (бир кунда)
    # Қатнашчилар:
    #  26: Жами  27: 18гача(офф)  28: 18катта(офф)  29: 18гача(онл)  30: 18катта(онл)
    # ОАВ:
    #  31: ТВ  32: Радио  33: Газета  34: Интернет  35: Видео жами  36:10К  37:100К  38:1М
    # Материаллар:
    #  39: Ижт.тармоқ  40: Ўз ташаб.  41: Флаер  42: LED  43: Бошқа
    # Пробация: 44
    # Қўшимча: 45:ИИО хизмат  46:Ҳамкор  47:Сайбер
    # ЖАМИ: 48

    NC = 48

    # ── Qator 1: Asosiy sarlavha ──────────────────────────────────────────────
    ws.merge_cells(f'A1:{get_column_letter(NC)}1')
    c1 = ws.cell(1, 1, f"«Хавфсиз ва Соғлом Юрт» — Кунлик қилинган ишлар ҳисоботи ({period})")
    c1.font = fnt(bold=True, color='FFFFFF', size=12)
    c1.fill = HDR1; c1.alignment = ctr
    ws.row_dimensions[1].height = 28

    # ── Qator 2: Guruh sarlavhalari ───────────────────────────────────────────
    h(2, 1, '№',              HDR1, fc='FFFFFF', rows=2)
    h(2, 2, 'Ҳудудлар',       HDR1, fc='FFFFFF', rows=2)
    h(2, 3, 'Ўтказилган учрашувлар сони', HDR2, fc='000000', span=23)
    h(2, 26,'Қатнашчилар сони',           QATN, fc='000000', span=5)
    h(2, 31,'ОАВ чиқишлари сони',         OAV_H, fc='000000', span=8)
    h(2, 39,'Тарқатилган материаллар',     MAT_H, fc='000000', span=5)
    h(2, 44,'Пробация',                    PROB,  fc='000000')
    h(2, 45,"Қўшимча кўрсаткичлар",       YELLOW,fc='000000', span=3)
    h(2, 48,'ЖАМИ учр.+проб.',             JAMI_F,fc='FFFFFF')
    ws.row_dimensions[2].height = 30

    # ── Qator 3: Kategoriya sarlavhalari ──────────────────────────────────────
    # Uchrashuvlar — har biri 2 col (сони | бир кунда)
    uch_cats = [
        'Маҳаллалар',       'Қизил МФЙлар',   'Таълим\nмуассасалари',
        'Касалхона',        'Бозорлар',        'Истироҳат\nбоғлари',
        'Жамоат\nтранспорти','Масжидлар',       'ҲМҚО',
        'Бошқа\n(Telegram)',
    ]
    for i, nom in enumerate(uch_cats):
        col = 3 + i * 2
        h(3, col, nom, HDR2, fc='000000', span=2, sz=8)

    h(3, 23, 'ИИО ТВ\nмурожаати',  GREEN, fc='FF0000', span=2, sz=8)
    h(3, 25, 'ЖАМИ\nучрашувлар',   YELLOW, fc='000000', sz=8)

    # Qatnashchilar
    for i, nom in enumerate(['ЖАМИ\nфуқаролар','18 ёшгача\n(офф.)','18 катта\n(офф.)','18 ёшгача\n(онл.)','18 катта\n(онл.)']):
        h(3, 26+i, nom, QATN, fc='000000', sz=8)

    # OAV
    for i, nom in enumerate(['Теле-\nвидение','Радио','Газета+\nЖурнал','Интернет','Видео\nжами','10К+','100К+','1М+']):
        h(3, 31+i, nom, OAV_H, fc='000000', sz=8)

    # Materiallar
    for i, nom in enumerate(['Ижт.\nтармоқ','Ўз\nташаббуси','Флаер+\nбукл.','LED\nэкран','Бошқа\nмат.']):
        h(3, 39+i, nom, MAT_H, fc='000000', sz=8)

    h(3, 44, 'Пробация\nсони',    PROB,  fc='000000', sz=8)
    h(3, 45, 'ИИО\nхизмат',       YELLOW,fc='000000', sz=8)
    h(3, 46, 'Ҳамкор\nташкилот',  YELLOW,fc='000000', sz=8)
    h(3, 47, 'Сайбер\nжиноят',    YELLOW,fc='000000', sz=8)
    h(3, 48, 'ЖАМИ\nучр.+проб.',  JAMI_F,fc='FFFFFF', sz=8)
    ws.row_dimensions[3].height = 42

    # ── Qator 4: сони / бир кунда subheader ──────────────────────────────────
    h(4, 1, '',   GRAY)
    h(4, 2, '',   GRAY)
    for i in range(10):
        col = 3 + i * 2
        h(4, col,   'сони',     SONI, bold=False, fc='000000', sz=8)
        h(4, col+1, 'бир кунда', BIRK, bold=False, fc='000000', sz=8)
    h(4, 23, 'ой\nбошидан', GREEN,  bold=True, fc='C00000', sz=8)
    h(4, 24, 'бир кунда',   BIRK,   bold=True, fc='C00000', sz=8)
    h(4, 25, 'бир кунда',   YELLOW, bold=True, fc='C00000', sz=8)
    for col in range(26, NC+1):
        fill = QATN if col<=30 else (OAV_H if col<=38 else (MAT_H if col<=43 else (PROB if col==44 else (YELLOW if col<=47 else JAMI_F))))
        fc   = '000000' if col < 48 else 'FFFFFF'
        h(4, col, 'бир кунда' if col!=48 else '', fill, bold=False, fc=fc, sz=8)
    ws.row_dimensions[4].height = 28

    # ── Ustun kengliklari ─────────────────────────────────────────────────────
    ws.column_dimensions['A'].width = 4
    ws.column_dimensions['B'].width = 24
    for col in range(3, NC+1):
        ws.column_dimensions[get_column_letter(col)].width = 6.5

    # ── Ma'lumot qatorlari ────────────────────────────────────────────────────
    def iv(d, key): return int(d.get(key) or 0) if d else 0
    def ginf(inf, k): return int(getattr(inf, k, 0) or 0) if inf else 0

    jami = {c: 0 for c in range(3, NC+1)}

    for ri, v in enumerate(viloyatlar, 5):
        b   = bot.get(v.id, {})
        kd  = ki.get(v.id, {})
        inf = infra.get(v.id)
        mfy = mfy_cnt.get(v.id, 0)

        talim_soni = ginf(inf,'oliy_talim')+ginf(inf,'akademik_litsey')+ginf(inf,'orta_talim')+ginf(inf,'maktabgacha')
        talim_birk = iv(b,'k2')+iv(b,'k3')+iv(b,'k4')+iv(b,'k5')
        jami_birk  = sum(iv(b,f'k{i}') for i in range(1,13))

        # сони | бир кунда — 10 kategoriya × 2 col
        pairs = [
            (mfy,                      iv(b,'k1')),   # Mahalla
            (ginf(inf,'qizil_mfy'),    iv(b,'k1')),   # Qizil MFY
            (talim_soni,               talim_birk),   # Ta'lim
            (ginf(inf,'kasalxona'),    iv(b,'k6')),   # Kasalxona
            (ginf(inf,'bozor'),        iv(b,'k7')),   # Bozor
            (ginf(inf,'istirohat'),    iv(b,'k10')),  # Istirohat
            (ginf(inf,'jamoat_transport'), iv(b,'k11')),  # Transport
            (ginf(inf,'masjid'),       iv(b,'k12')),  # Masjid
            (ginf(inf,'xmko'),         iv(b,'k8')),   # XMKO
            (ginf(inf,'telegram'),     iv(b,'k9')),   # Boshqa/Telegram
        ]

        row_vals = {}
        for i, (soni, birk) in enumerate(pairs):
            row_vals[3 + i*2]   = soni
            row_vals[3 + i*2+1] = birk

        iiotv_bk = iv(kd,'iiotv')
        row_vals[23] = 0          # IIO TV "oy boshidan" — bizda yo'q
        row_vals[24] = iiotv_bk   # IIO TV bir kunda
        row_vals[25] = jami_birk  # JAMI uchrashuvlar

        row_vals[26] = iv(b,'fuk')
        row_vals[27] = iv(b,'o18g'); row_vals[28] = iv(b,'o18k')
        row_vals[29] = iv(b,'n18g'); row_vals[30] = iv(b,'n18k')

        row_vals[31] = iv(kd,'tv');    row_vals[32] = iv(kd,'radio')
        row_vals[33] = iv(kd,'gazeta'); row_vals[34] = iv(kd,'internet')
        row_vals[35] = iv(kd,'video')
        row_vals[36] = iv(kd,'v10k');  row_vals[37] = iv(kd,'v100k'); row_vals[38] = iv(kd,'v1m')

        row_vals[39] = iv(kd,'ijt');  row_vals[40] = iv(kd,'ozt')
        row_vals[41] = iv(kd,'flay'); row_vals[42] = iv(kd,'led'); row_vals[43] = iv(kd,'matb')

        row_vals[44] = iv(b,'k13')
        row_vals[45] = iv(kd,'iiox'); row_vals[46] = iv(kd,'hmkr'); row_vals[47] = iv(kd,'sayb')
        row_vals[48] = jami_birk + iv(b,'k13')

        # Raqam va hudud
        cc = ws.cell(ri, 1, ri-4)
        cc.font = fnt(); cc.fill = GRAY; cc.border = brd; cc.alignment = ctr

        cc = ws.cell(ri, 2, v.nomi)
        cc.font = fnt(bold=True); cc.fill = GRAY; cc.border = brd; cc.alignment = lft

        # Har ustun
        for col, val in row_vals.items():
            # "сони" ustunlari: 3,5,7,9,11,13,15,17,19,21 (toq)
            is_soni = (3 <= col <= 22) and (col % 2 == 1)
            is_birk = (3 <= col <= 22) and (col % 2 == 0)
            if col == 23: fill = GREEN
            elif col == 24: fill = BIRK
            elif col == 25: fill = YELLOW
            elif col <= 30: fill = QATN
            elif col <= 38: fill = OAV_D
            elif col <= 43: fill = MAT_D
            elif col == 44: fill = PROB
            elif col <= 47: fill = PatternFill('solid', fgColor='FFFFCC')
            elif col == 48: fill = YELLOW
            elif is_soni: fill = SONI
            else: fill = BIRK

            cc = ws.cell(ri, col, val)
            cc.font = fnt()
            cc.fill = fill; cc.border = brd; cc.alignment = ctr
            jami[col] = jami.get(col, 0) + val

        ws.row_dimensions[ri].height = 15

    # ── ЖАМИ qatori ───────────────────────────────────────────────────────────
    jr = len(viloyatlar) + 5
    ws.cell(jr, 1).value = ''
    cc = ws.cell(jr, 2, 'ЖАМИ')
    cc.font = fnt(bold=True, color='FFFFFF', size=10); cc.fill = JAMI_F; cc.border = brd; cc.alignment = ctr
    for col in range(3, NC+1):
        cc = ws.cell(jr, col, jami.get(col, 0))
        cc.font = fnt(bold=True, color='FFFFFF', size=9)
        cc.fill = JAMI_F; cc.border = brd; cc.alignment = ctr
    for col in (1,):
        cc = ws.cell(jr, col)
        cc.fill = JAMI_F; cc.border = brd
    ws.row_dimensions[jr].height = 18

    ws.freeze_panes = 'C5'

    buf = _io.BytesIO()
    wb.save(buf); buf.seek(0)
    fname = f"kunlik_ishlar_{start_str}_{end_str}.xlsx"
    return HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'}
    )


# ── Murojaat ──────────────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def murojaat_usullar(request):
    return Response(MurojaatUsulSerializer(MurojaatUsul.objects.all(), many=True).data)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def murojaat_kasblar(request):
    return Response(MurojaatKasbSerializer(MurojaatKasb.objects.all(), many=True).data)


@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def murojaat_list_create(request):
    vf = get_viloyat_qs_filter(request, 'viloyat_id')

    if request.method == 'GET':
        start  = request.GET.get('start')
        end    = request.GET.get('end')
        tuman  = request.GET.get('tuman_id')
        qs = Murojaat.objects.filter(**vf).select_related('viloyat', 'tuman', 'mahalla', 'usul', 'kasb')
        if start:
            qs = qs.filter(sana__gte=start)
        if end:
            qs = qs.filter(sana__lte=end)
        if tuman:
            qs = qs.filter(tuman_id=tuman)
        return Response(MurojaatSerializer(qs[:500], many=True).data)

    # POST — yangi murojaat qo'shish
    data = request.data.copy()
    if request.user.role == 'viloyat':
        data['viloyat'] = request.user.viloyat_id
    ser = MurojaatSerializer(data=data)
    if not ser.is_valid():
        return Response(ser.errors, status=400)
    obj = ser.save(yaratuvchi=request.user)
    audit(request, 'murojaat_qoshish', f"ID={obj.id} sana={obj.sana}")

    # ── Mahalla inspektoriga Telegram xabar yuborish ──────────────────────────
    if obj.mahalla_id:
        try:
            m = Mahalla.objects.get(pk=obj.mahalla_id)
            tg_id = m.tg_id  # inspektor Telegram ID
            if tg_id and tg_id != 0:
                fish    = obj.fish or "F.I.O. ko'rsatilmagan"
                telefon = obj.telefon or "ko'rsatilmagan"
                matn = (
                    f"📢 <b>Profilaktik suhbat o'tkazish to'g'risida xabarnoma</b>\n\n"
                    f"Sizning mahallangizda yashovchi <b>{fish}</b> kiberfiribgarlarga aldanganligini ma'lum qilamiz.\n\n"
                    f"Ushbu fuqaro bilan profilaktik suhbat o'tkazib, botga targ'ibotingizni joylashtiring.\n\n"
                    f"📞 <b>Telefon:</b> {telefon}"
                )
                tg_send(tg_id, matn)
        except Exception:
            pass  # Xabar yuborilmasa ham murojaat saqlanadi

    return Response(MurojaatSerializer(obj).data, status=201)


@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def murojaat_detail(request, pk):
    try:
        obj = Murojaat.objects.get(pk=pk)
    except Murojaat.DoesNotExist:
        return Response({'error': 'Topilmadi'}, status=404)
    # Viloyat admin faqat o'z viloyatini ko'ra/o'zgartira oladi
    if request.user.role == 'viloyat' and obj.viloyat_id != request.user.viloyat_id:
        return Response({'error': 'Ruxsat yo\'q'}, status=403)

    if request.method == 'GET':
        return Response(MurojaatSerializer(obj).data)

    if request.method == 'PUT':
        ser = MurojaatSerializer(obj, data=request.data, partial=True)
        if not ser.is_valid():
            return Response(ser.errors, status=400)
        ser.save()
        audit(request, 'murojaat_yangilash', f"ID={obj.id}")
        return Response(MurojaatSerializer(obj).data)

    # DELETE
    audit(request, 'murojaat_ochirish', f"ID={obj.id} sana={obj.sana}")
    obj.delete()
    return Response({'ok': True})


# ── Murojaat import ───────────────────────────────────────────────────────────
MUROJAAT_IMPORT_HEADERS = [
    'Sana (YYYY-MM-DD)', 'Viloyat', 'Tuman', 'Mahalla', 'F.I.SH', 'Jinsi (erkak/ayol)',
    'Yoshi', 'Telefon', 'Sodir etish usuli', 'Kasbi', 'Kasb izohi',
    "O'quv muassasasi", 'Kurs', 'Ijtimoiy tarmoq', 'Zarar (so\'m)',
    'Holat (yangi/takroriy/aybi/togri)', 'Fabula',
]

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def murojaat_shablon(request):
    """GET /api/murojaat/shablon/ — import uchun bo'sh Excel shablon."""
    misol = [
        '2026-01-15', 'Toshkent viloyati', 'Chirchiq', 'Guliston MFY',
        'Aliyev Vali Aliyevich', 'erkak', 35, '+998901234567',
        'Telegram orqali firibgarlik', "Ishchi", '', '', '', 'telegram',
        1500000, 'yangi', "Qisqacha voqea bayoni...",
    ]
    audit(request, 'murojaat_shablon', 'Import shabloni yuklab olindi')
    return excel_response(MUROJAAT_IMPORT_HEADERS, [misol], 'murojaat_shablon.xlsx')


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def murojaat_import(request):
    """POST /api/murojaat/import/ — to'ldirilgan Excel faylni bazaga import qiladi."""
    f = request.FILES.get('fayl')
    if not f:
        return Response({'error': 'Excel fayl yuklanmadi'}, status=400)

    role = request.user.role
    if role not in ('viloyat', 'respublika'):
        return Response({'error': 'Ruxsat yo\'q'}, status=403)

    try:
        wb = openpyxl.load_workbook(f, data_only=True)
        ws = wb.active
    except Exception:
        return Response({'error': 'Excel faylni o\'qib bo\'lmadi'}, status=400)

    JINSI_MAP  = {'erkak': 'erkak', 'ayol': 'ayol', 'erkak ': 'erkak', 'ayol ': 'ayol'}
    HOLAT_MAP  = {'yangi': 'yangi', 'takroriy': 'takroriy', 'aybi': 'aybi', 'togri': 'togri'}
    TARMOQ_MAP = {'telegram': 'telegram', 'instagram': 'instagram', 'facebook': 'facebook',
                  'tiktok': 'tiktok', 'bigolive': 'bigolive', 'boshqa': 'boshqa'}

    usullar = list(MurojaatUsul.objects.all())
    kasblar = list(MurojaatKasb.objects.all())

    def topish(nomi, ro_yxat):
        if not nomi:
            return None
        nomi = str(nomi).strip().lower()
        for item in ro_yxat:
            if item.nomi.strip().lower() == nomi:
                return item
        for item in ro_yxat:
            if nomi in item.nomi.strip().lower():
                return item
        return None

    created = 0
    errors  = []
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    for idx, row in enumerate(rows, start=2):
        if not row or not any(row):
            continue
        (sana, viloyat_nomi, tuman_nomi, mahalla_nomi, fish, jinsi, yosh, telefon,
         usul_nomi, kasb_nomi, kasb_izoh, kasb_muassasa, kasb_kurs,
         tarmoq, zarar, holat, fabula) = (list(row) + [None] * 17)[:17]

        try:
            if not sana:
                errors.append({'qator': idx, 'sabab': 'Sana kiritilmagan'})
                continue
            sana_val = sana.date() if hasattr(sana, 'date') else datetime.strptime(str(sana), '%Y-%m-%d').date()

            if role == 'viloyat':
                viloyat_id = request.user.viloyat_id
            else:
                vil = Viloyat.objects.filter(nomi__icontains=str(viloyat_nomi or '').strip()).first()
                if not vil:
                    errors.append({'qator': idx, 'sabab': f'Viloyat topilmadi: {viloyat_nomi}'})
                    continue
                viloyat_id = vil.id

            tuman = Tuman.objects.filter(
                viloyat_id=viloyat_id, tuman_nomi__icontains=str(tuman_nomi or '').strip()
            ).first()
            if not tuman:
                errors.append({'qator': idx, 'sabab': f'Tuman topilmadi: {tuman_nomi}'})
                continue

            mahalla = None
            if mahalla_nomi:
                mahalla = Mahalla.objects.filter(
                    tuman_id=tuman.id, mahalla_nomi__icontains=str(mahalla_nomi).strip()
                ).first()

            usul = topish(usul_nomi, usullar)
            kasb = topish(kasb_nomi, kasblar)

            Murojaat.objects.create(
                sana=sana_val,
                viloyat_id=viloyat_id,
                tuman_id=tuman.id,
                mahalla_id=mahalla.id if mahalla else None,
                fish=str(fish or '').strip(),
                jinsi=JINSI_MAP.get(str(jinsi or '').strip().lower(), ''),
                telefon=str(telefon or '').strip(),
                fabula=str(fabula or '').strip(),
                zarar=zarar or None,
                usul=usul,
                kasb=kasb,
                yosh=int(yosh) if yosh else None,
                holat=HOLAT_MAP.get(str(holat or '').strip().lower(), 'yangi'),
                ijtimoiy_tarmoq=TARMOQ_MAP.get(str(tarmoq or '').strip().lower(), ''),
                kasb_izoh=str(kasb_izoh or '').strip(),
                kasb_muassasa=str(kasb_muassasa or '').strip(),
                kasb_kurs=int(kasb_kurs) if kasb_kurs else None,
                yaratuvchi=request.user,
            )
            created += 1
        except Exception as e:
            errors.append({'qator': idx, 'sabab': str(e)})

    audit(request, 'murojaat_import', f"{created} ta import qilindi, {len(errors)} ta xato")
    return Response({'created': created, 'errors': errors})


# ── Murojaat Hisobot ──────────────────────────────────────────────────────────
def _cnt(qs, v_ids):
    res = {vid: 0 for vid in v_ids}
    res['total'] = qs.count()
    for r in qs.values('viloyat_id').annotate(n=Count('id')):
        if r['viloyat_id'] in res:
            res[r['viloyat_id']] = r['n']
    return res


def _build_hisobot_rows(start, end, v_ids):
    """Barcha hisobot qatorlarini hisoblash (API va Excel uchun umumiy)"""
    today_str = date.today().isoformat()
    P = Murojaat.objects.filter(sana__gte=start, sana__lte=end)
    T = Murojaat.objects.filter(sana=today_str)

    def kf(ids):
        return lambda q: q.filter(kasb_id__in=ids)
    def hf(holat):
        return lambda q: q.filter(holat=holat)
    def jf(jinsi, ymin=None, ymax=None):
        def fn(q):
            q = q.filter(jinsi=jinsi)
            if ymin: q = q.filter(yosh__gte=ymin)
            if ymax: q = q.filter(yosh__lte=ymax)
            return q
        return fn

    def r(tartib, nomi, daraja, f):
        return {'tartib': tartib, 'nomi': nomi, 'daraja': daraja,
                'jami': _cnt(f(P), v_ids), 'bugun': _cnt(f(T), v_ids)}

    af = lambda q: q

    rows = [
        {'section': 'ИЖТИМОИЙ ҲОЛАТИ'},
        r('ЖАМИ', 'Аҳоли сони жами',                                   0, af),
        r('1',    'Хусусий сектор тадбиркорлар',                       1, kf([1,10,11,12,13])),
        r('1.1',  'Савдо фаолияти билан шуғулланувчилар',              2, kf([10])),
        r('1.2',  'Йирик корхона ишчи-хизматчилари',                   2, kf([11])),
        r('1.3',  'Кичик ташкилотлар ишчи-хизматчилари',               2, kf([12])),
        r('1.4',  'Хизмат кўрсатиш сектори ишчилари',                  2, kf([13])),
        r('2',    'Тўлов ташкилотлари',                                 1, kf([2,14,15])),
        r('2.1',  'Банклар (Ипотека, Асака, Капитал ва бошқалар)',      2, kf([14])),
        r('2.2',  'Тўлов ташкилотлари (Payme, Click, Uzum, Upay...)',  2, kf([15])),
        r('3',    'Таълим ташкилотлари',                                1, kf([3,16,17,18,19,20,21,22,23,24,25])),
        r('3.1',  'Олий таълим муассасалари (Университет, Институт)',  2, kf([16,20,21])),
        r('',     '  Ўқитувчи (таълим соҳасидаги ходимлар)',           3, kf([20])),
        r('',     '  Талабалар',                                         3, kf([21])),
        r('3.2',  'Ўрта-махсус муассасалари (Академик лицей, Техникум)', 2, kf([17,22,23])),
        r('',     '  Ўқитувчи (таълим соҳасидаги ходимлар)',           3, kf([22])),
        r('',     '  Талабалар',                                         3, kf([23])),
        r('3.3',  'Ўрта таълим муассасалари (мактаб, ихтисослаштирилган)', 2, kf([18,24,25])),
        r('',     '  Ўқитувчи (таълим соҳасидаги ходимлар)',           3, kf([24])),
        r('',     '  Мактаб ўқувчилари',                                3, kf([25])),
        r('3.4',  'Мактабгача таълим муассасалари',                     2, kf([19])),
        r('4',    'Тиббиёт соҳасидаги ходимлар',                       1, kf([4])),
        r('5',    'Давлат корхона ва ташкилотлар ишчилари',             1, kf([5])),
        r('6',    'ҲМҚО органлари ходимлари',                          1, kf([6])),
        r('7',    'Нафақахўрлар',                                        1, kf([7])),
        r('8',    'Вақтинча ишсизлар',                                  1, kf([8])),
        r('9',    'Чет эл фуқаролари',                                  1, kf([9])),
        r('10',   'Такрорий мурожатлар',                                1, hf('takroriy')),
        r('11',   'Фуқаронинг ўз айби билан',                           1, hf('aybi')),
        r('12',   'Тўғридан тўғри ариза қабул қилиниши',               1, hf('togri')),

        {'section': 'ЖАБРЛАНУВЧИНИНГ ЁШИ ВА ЖИНСИ КЕСИМИДА'},
        r('11',   'Жинси ва ёши жами',                                  0, af),
        r('',     'ЭРКАК',                                               1, jf('erkak')),
        r('',     '  18 ёшгача',                                         2, jf('erkak', None, 17)),
        r('',     '  18 дан 30 ёшгача',                                  2, jf('erkak', 18, 30)),
        r('',     '  31 ёшдан 40 ёшгача',                                2, jf('erkak', 31, 40)),
        r('',     '  41 ёшдан 50 ёшгача',                                2, jf('erkak', 41, 50)),
        r('',     '  51 ёшдан 60 ёшгача',                                2, jf('erkak', 51, 60)),
        r('',     '  61 ёшдан ошган',                                    2, jf('erkak', 61, None)),
        r('',     'АЁЛ',                                                 1, jf('ayol')),
        r('',     '  18 ёшгача',                                         2, jf('ayol', None, 17)),
        r('',     '  18 дан 30 ёшгача',                                  2, jf('ayol', 18, 30)),
        r('',     '  31 ёшдан 40 ёшгача',                                2, jf('ayol', 31, 40)),
        r('',     '  41 ёшдан 50 ёшгача',                                2, jf('ayol', 41, 50)),
        r('',     '  51 ёшдан 60 ёшгача',                                2, jf('ayol', 51, 60)),
        r('',     '  61 ёшдан ошган',                                    2, jf('ayol', 61, None)),

        {'section': 'СОДИР ЭТИШ УСУЛИГА КЎРА'},
        r('ЖАМИ', 'Усуллар жами',                                       0, af),
    ]

    usullar = MurojaatUsul.objects.filter(ota_id__isnull=True).order_by('tartib')
    for i, u in enumerate(usullar, 12):
        children = list(MurojaatUsul.objects.filter(ota_id=u.id).values_list('id', flat=True))
        if children:
            all_ids = [u.id] + children
            rows.append(r(str(i), u.nomi, 1, lambda q, ids=all_ids: q.filter(usul_id__in=ids)))
            for c in MurojaatUsul.objects.filter(ota_id=u.id).order_by('tartib'):
                rows.append(r('', '  ' + c.nomi, 2, lambda q, cid=c.id: q.filter(usul_id=cid)))
        else:
            rows.append(r(str(i), u.nomi, 1, lambda q, uid=u.id: q.filter(usul_id=uid)))

    return rows


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def murojaat_hisobot(request):
    start = request.GET.get('start', date.today().replace(day=1).isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    viloyatlar = list(Viloyat.objects.values('id', 'nomi').order_by('id'))
    v_ids = [v['id'] for v in viloyatlar]
    rows = _build_hisobot_rows(start, end, v_ids)
    return Response({'viloyatlar': viloyatlar, 'rows': rows, 'start': start, 'end': end})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def murojaat_hisobot_excel(request):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import io

    start = request.GET.get('start', date.today().replace(day=1).isoformat())
    end   = request.GET.get('end',   date.today().isoformat())

    viloyatlar = list(Viloyat.objects.values('id', 'nomi').order_by('id'))
    v_ids = [v['id'] for v in viloyatlar]
    rows  = _build_hisobot_rows(start, end, v_ids)

    wb  = Workbook()
    ws  = wb.active
    ws.title = end[:10]

    # ── Stil konstantalari ──
    thin   = Side(style='thin',   color='AAAAAA')
    medium = Side(style='medium', color='1F4E79')
    brd    = Border(left=thin, right=thin, top=thin, bottom=thin)
    mbrd   = Border(left=medium, right=medium, top=medium, bottom=medium)
    ctr    = Alignment(horizontal='center', vertical='center', wrap_text=True)
    lft    = Alignment(horizontal='left',   vertical='center', wrap_text=True)

    def fill(hex_): return PatternFill('solid', start_color=hex_)
    def fnt(bold=False, color='000000', size=9, italic=False):
        return Font(name='Times New Roman', bold=bold, color=color, size=size, italic=italic)

    DARK_BLUE = fill('1F4E79')
    MID_BLUE  = fill('2E75B6')
    DARK_GRN  = fill('375623')
    LT_GREEN  = fill('E2EFDA')
    LT_BLUE   = fill('DEEAF1')
    LT_GRAY   = fill('F2F2F2')
    WHITE_F   = fill('FFFFFF')

    v_count   = len(viloyatlar)
    # Ustunlar: A(№), B(Ko'rsatkich), C(JAMI), D(Bir kunda),
    # E,F (viloyat1_jami, viloyat1_bugun), G,H, ...
    total_cols = 2 + 1 + 1 + v_count * 2   # №+nomi + jami+bugun + viloyatlar

    # ── 1-qator: Sarlavha ──
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=total_cols)
    c = ws.cell(1, 1, f'KIBERJINOYAT MUROJAATLARI HISOBOTI   ({start} — {end})')
    c.font      = fnt(bold=True, color='FFFFFF', size=12)
    c.fill      = DARK_BLUE
    c.alignment = ctr
    c.border    = mbrd
    ws.row_dimensions[1].height = 35

    # ── 2-qator: Ustun nomlari (1-qator) ──
    col_names = ['№', 'Кўрсаткичлар', 'ЖАМИ', 'Бир кунда']
    for v in viloyatlar:
        short = v['nomi'].replace(' viloyati','').replace(' viloayti','')
        col_names += [short, 'Бир кунда']
    for ci, h in enumerate(col_names, 1):
        c = ws.cell(2, ci, h)
        c.font = fnt(bold=True, color='FFFFFF', size=9)
        c.fill = MID_BLUE
        c.alignment = ctr
        c.border = brd
    ws.row_dimensions[2].height = 45

    # ── Ustun kengliklari ──
    ws.column_dimensions['A'].width = 6
    ws.column_dimensions['B'].width = 52
    ws.column_dimensions['C'].width = 9
    ws.column_dimensions['D'].width = 9
    for ci in range(v_count):
        ws.column_dimensions[get_column_letter(5 + ci*2)].width = 10
        ws.column_dimensions[get_column_letter(6 + ci*2)].width = 9

    # ── Ma'lumot qatorlari ──
    er = 3  # excel row counter
    for row in rows:
        # SECTION HEADER
        if 'section' in row:
            ws.merge_cells(start_row=er, start_column=1, end_row=er, end_column=total_cols)
            c = ws.cell(er, 1, row['section'])
            c.font = fnt(bold=True, color='FFFFFF', size=10)
            c.fill = DARK_BLUE
            c.alignment = ctr
            c.border = brd
            ws.row_dimensions[er].height = 22
            er += 1
            continue

        d  = row.get('daraja', 0)
        jd = row.get('jami',   {})
        bd = row.get('bugun',  {})

        # Qator fonlari
        if d == 0:
            rfill = LT_GRAY;  bfill = fill('E8E8E8'); tf = fnt(bold=True)
        elif d == 1:
            rfill = DARK_GRN; bfill = LT_GREEN;       tf = fnt(bold=True, color='FFFFFF')
        elif d == 2:
            rfill = LT_BLUE;  bfill = fill('EAF4FB');  tf = fnt()
        else:
            rfill = WHITE_F;  bfill = LT_GRAY;         tf = fnt(italic=True, color='444444', size=8)

        # № ustuni
        c = ws.cell(er, 1, row.get('tartib', ''))
        c.font = tf; c.fill = rfill; c.alignment = ctr; c.border = brd

        # Nomi
        c = ws.cell(er, 2, row['nomi'])
        c.font = tf; c.fill = rfill; c.alignment = lft; c.border = brd

        # JAMI (davr)
        c = ws.cell(er, 3, jd.get('total', 0) or 0)
        c.font = fnt(bold=(d<=1), color='FFFFFF' if d==1 else '000000')
        c.fill = rfill; c.alignment = ctr; c.border = brd

        # Bir kunda
        c = ws.cell(er, 4, bd.get('total', 0) or 0)
        c.font = fnt(italic=True, color='555555', size=8)
        c.fill = bfill; c.alignment = ctr; c.border = brd

        # Viloyatlar
        for ci, v in enumerate(viloyatlar):
            vj = ws.cell(er, 5+ci*2,   jd.get(v['id'], 0) or 0)
            vb = ws.cell(er, 6+ci*2,   bd.get(v['id'], 0) or 0)
            vj.font = fnt(); vj.fill = rfill if d==0 else WHITE_F
            vj.alignment = ctr; vj.border = brd
            vb.font = fnt(italic=True, color='666666', size=8)
            vb.fill = bfill; vb.alignment = ctr; vb.border = brd

        ws.row_dimensions[er].height = 18
        er += 1

        # "Бир кунда" alohida pastki qator
        ws.cell(er, 1).fill = bfill; ws.cell(er, 1).border = brd
        c2 = ws.cell(er, 2, '    Бир кунда')
        c2.font = fnt(italic=True, color='666666', size=8)
        c2.fill = bfill; c2.alignment = lft; c2.border = brd
        c3 = ws.cell(er, 3, bd.get('total', 0) or 0)
        c3.font = fnt(italic=True, color='555555', size=8)
        c3.fill = bfill; c3.alignment = ctr; c3.border = brd
        c4 = ws.cell(er, 4, '')
        c4.fill = bfill; c4.border = brd
        for ci, v in enumerate(viloyatlar):
            vb = ws.cell(er, 5+ci*2, bd.get(v['id'], 0) or 0)
            vb.font = fnt(italic=True, color='666666', size=8)
            vb.fill = bfill; vb.alignment = ctr; vb.border = brd
            ws.cell(er, 6+ci*2).fill = bfill; ws.cell(er, 6+ci*2).border = brd
        ws.row_dimensions[er].height = 13
        er += 1

    ws.freeze_panes = 'C3'

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f'murojaat_hisobot_{start}_{end}.xlsx'
    return HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename="{fname}"'}
    )


# ════════════════════════════════════════════════════════════════════
#  KUNLIK ISHLAR
# ════════════════════════════════════════════════════════════════════

def _bot_agg(viloyat_id, sana_str):
    """Bot hisobotlaridan shu viloyat + sana uchun agregatsiya."""
    from django.db.models import Sum, Count, Q
    sana = date.fromisoformat(sana_str)

    qs = Hisobot.objects.filter(
        status=2,
        mahalla__tuman__viloyat_id=viloyat_id,
        qushilgan_vaqt__date=sana,
        targibot_turi__in=[1, 2],
    )

    # kategoriya bo'yicha hisobotlar soni (bir kunda)
    kat_ids = {k: list(TargibotUtkazilganJoy.objects.filter(kategoriya=k).values_list('id', flat=True))
               for k in range(1, 12)}

    def c(k):
        return qs.filter(targibot_utgan_joy__in=kat_ids.get(k, [])).count()

    agg = qs.aggregate(
        fuk_jami    = Sum('qatnashchilar_soni'),
        off18g      = Sum('offline_18_gacha'),
        off18k      = Sum('offline_18_katta'),
        onl18g      = Sum('online_18_gacha'),
        onl18k      = Sum('online_18_katta'),
    )

    return {
        'kat1'     : c(1),   # Қизил МФЙлар
        'kat2'     : c(2),   # Таълим муассасалари
        'kat3'     : c(3),   # Касалхона ва поликлиника
        'kat4'     : c(4),   # Бозорлар ва йирик савдо мажмуалари
        'kat5'     : c(5),   # Истироҳат боғлари
        'kat6'     : c(6),   # Жамоат транспортлари
        'kat7'     : c(7),   # Масжидлар
        'kat8'     : c(8),   # ҲМҚО
        'kat9'     : c(9),   # Бошқа идора ва ташкилотлар
        'kat10'    : c(10),  # Аҳоли гавжум жойларда
        'kat11'    : c(11),  # Probatsiya ro'yxatidagi shaxslar
        'jami'     : qs.count(),
        'fuk_jami' : int(agg['fuk_jami'] or 0),
        'off18g'  : int(agg['off18g'] or 0),
        'off18k'  : int(agg['off18k'] or 0),
        'onl18g'  : int(agg['onl18g'] or 0),
        'onl18k'  : int(agg['onl18k'] or 0),
    }


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def kunlik_ishlar_get(request):
    """GET /api/kunlik-ishlar/?viloyat=&sana="""
    role = request.user.role

    if role == 'viloyat':
        viloyat_id = request.user.viloyat_id
    else:
        viloyat_id = request.GET.get('viloyat')
    sana = request.GET.get('sana', date.today().isoformat())

    if not viloyat_id:
        return Response({'error': 'viloyat kerak'}, status=400)

    record, _ = KunlikIshlar.objects.get_or_create(
        viloyat_id=viloyat_id, sana=sana
    )
    bot_data = _bot_agg(viloyat_id, sana)

    kat_nomlar = {f'kat{k}': v for k, v in TargibotUtkazilganJoy.KATEGORIYA}

    return Response({
        'record'    : KunlikIshlarSerializer(record).data,
        'bot'       : bot_data,
        'kat_nomlar': kat_nomlar,
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def kunlik_ishlar_save(request):
    """POST /api/kunlik-ishlar/saqlash/ — viloyat saqlaydi/yuboradi"""
    role = request.user.role
    if role not in ('viloyat', 'respublika'):
        return Response({'error': 'Ruxsat yo\'q'}, status=403)

    viloyat_id = request.user.viloyat_id if role == 'viloyat' else request.data.get('viloyat')
    sana       = request.data.get('sana', date.today().isoformat())

    record, _ = KunlikIshlar.objects.get_or_create(viloyat_id=viloyat_id, sana=sana)

    if record.status == 3:
        return Response({'error': 'Tasdiqlangan hisobot o\'zgartirib bo\'lmaydi'}, status=400)

    FIELDS = [
        'qizil_mfy_soni','istirohat_soni','transport_soni','masjid_soni','iio_tv_murojaati',
        'uchrashuv_proof_url','uchrashuv_proof_rasm',
        'oav_tv_soni','oav_tv_url','oav_radio_soni','oav_radio_url',
        'oav_gazeta_jurnal_soni','oav_gazeta_jurnal_url',
        'oav_video_soni','oav_video_10k','oav_video_100k','oav_video_1m','oav_video_url',
        'oav_internet_soni','oav_internet_url',
        'mat_ijtimoiy_tarmoq','mat_oz_tashabbusi','mat_flayer_buklet',
        'mat_led_ekran','mat_boshqa','mat_proof_url','mat_proof_rasm',
        'suhbat_soni','suhbat_proof_url','suhbat_proof_rasm',
        'iio_xizmat_soni','hamkor_tashkilot_soni','sayber_soni',
    ]
    for f in FIELDS:
        if f in request.data:
            setattr(record, f, request.data[f])

    yuborish = request.data.get('yuborish', False)
    if yuborish:
        record.status = 2
        record.rad_sababi = ''

    record.save()
    audit(request, 'kunlik_ishlar_saqlash', f"{viloyat_id} | {sana}")
    return Response(KunlikIshlarSerializer(record).data)


@api_view(['POST'])
@permission_classes([IsAuthenticated, IsRespublika])
def kunlik_ishlar_tasdiqlash(request):
    """POST /api/kunlik-ishlar/tasdiqlash/ — respublika tasdiqlaydi/rad etadi"""
    ids        = request.data.get('tasdiq_ids', [])
    rad_ids    = request.data.get('rad_ids', [])
    rad_sababi = request.data.get('rad_sababi', '')

    KunlikIshlar.objects.filter(id__in=ids).update(status=3, rad_sababi='')
    KunlikIshlar.objects.filter(id__in=rad_ids).update(status=4, rad_sababi=rad_sababi)

    audit(request, 'kunlik_ishlar_tasdiqlash',
          f"tasdiq:{ids} rad:{rad_ids}")
    return Response({'ok': True})


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def kunlik_ishlar_list(request):
    """GET /api/kunlik-ishlar/list/?sana= — respublika uchun barcha viloyatlar"""
    sana = request.GET.get('sana', date.today().isoformat())
    role = request.user.role

    if role == 'viloyat':
        qs = KunlikIshlar.objects.filter(viloyat_id=request.user.viloyat_id, sana=sana)
        result = []
        for rec in qs:
            d = KunlikIshlarSerializer(rec).data
            d['bot'] = _bot_agg(rec.viloyat_id, sana)
            result.append(d)
        return Response(result)

    # Respublika: barcha viloyatlarni ko'rsatamiz (record bo'lsa yoki bo'lmasa ham)
    viloyatlar = Viloyat.objects.all().order_by('nomi')
    records_by_viloyat = {
        r.viloyat_id: r
        for r in KunlikIshlar.objects.filter(sana=sana).select_related('viloyat')
    }

    result = []
    for v in viloyatlar:
        rec = records_by_viloyat.get(v.id)
        if rec:
            d = KunlikIshlarSerializer(rec).data
        else:
            d = {
                'id': None, 'viloyat': v.id, 'viloyat_nomi': v.nomi,
                'sana': sana, 'status': 1,
                'oav_tv_soni': 0, 'oav_radio_soni': 0, 'oav_gazeta_jurnal_soni': 0,
                'oav_internet_soni': 0, 'oav_video_soni': 0,
                'mat_ijtimoiy_tarmoq': 0, 'mat_flayer_buklet': 0,
                'hamkor_tashkilot_soni': 0, 'sayber_soni': 0, 'rad_sababi': '',
            }
        d['bot'] = _bot_agg(v.id, sana)
        result.append(d)

    return Response(result)


@api_view(['GET', 'PUT'])
@permission_classes([IsAuthenticated])
def infratuzilma(request):
    """GET/PUT /api/infratuzilma/ — viloyat targ'ibot joylar soni."""
    role = request.user.role
    if role == 'viloyat':
        viloyat_id = request.user.viloyat_id
    elif role == 'respublika':
        viloyat_id = request.GET.get('viloyat') or request.data.get('viloyat')
        if not viloyat_id:
            # Respublika: barcha viloyatlar uchun ro'yxat
            qs = ViloyatInfratuzilma.objects.select_related('viloyat').all()
            return Response(ViloyatInfratuzilmaSerializer(qs, many=True).data)
    else:
        return Response({'error': 'Ruxsat yo\'q'}, status=403)

    obj, _ = ViloyatInfratuzilma.objects.get_or_create(viloyat_id=viloyat_id)
    if request.method == 'GET':
        return Response(ViloyatInfratuzilmaSerializer(obj).data)

    ser = ViloyatInfratuzilmaSerializer(obj, data=request.data, partial=True)
    ser.is_valid(raise_exception=True)
    ser.save()
    return Response(ser.data)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def kunlik_ishlar_rasm(request):
    """POST /api/kunlik-ishlar/rasm/ — rasm yuklash"""
    import uuid as _uuid
    f = request.FILES.get('rasm')
    if not f:
        return Response({'error': 'Rasm kerak'}, status=400)
    ext      = f.name.rsplit('.', 1)[-1].lower()
    fname    = f"ki_{_uuid.uuid4().hex[:12]}.{ext}"
    fpath    = os.path.join(settings.MEDIA_ROOT, 'images', fname)
    with open(fpath, 'wb') as out:
        for chunk in f.chunks():
            out.write(chunk)
    return Response({'url': fname})


# ═══════════════════════════════════════════════════════════════════════════════
# TAHLILIY HISOBOTLAR
# ═══════════════════════════════════════════════════════════════════════════════

# ── 7. Samaradorlik hisoboti ──────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def samaradorlik_hisobot(request):
    """
    Respublika: viloyatlar bo'yicha targ'ibot vs murojaat.
    Viloyat: mahallalar bo'yicha targ'ibot vs murojaat.
    """
    start = request.GET.get('start', date(date.today().year, date.today().month, 1).isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    role  = request.user.role

    if role == 'respublika':
        sql = """
            SELECT v.nomi AS nomi,
                   (SELECT COUNT(*) FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
                    JOIN tuman t ON m.tuman_id=t.id
                    WHERE t.viloyat_id=v.id AND h.status=2 AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s) AS targibot_soni,
                   (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
                    JOIN tuman t ON m.tuman_id=t.id
                    WHERE t.viloyat_id=v.id AND h.status=2 AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s) AS qatnashchilar,
                   (SELECT COUNT(*) FROM murojaat mr WHERE mr.viloyat_id=v.id AND mr.sana BETWEEN %s AND %s) AS murojaat_soni
            FROM viloyat v ORDER BY v.nomi
        """
        params = [start, end] * 3
    else:
        viloyat_id = request.user.viloyat_id
        sql = """
            SELECT mahalla.mahalla_nomi AS nomi, tuman.tuman_nomi,
                   (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                    AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s) AS targibot_soni,
                   (SELECT COALESCE(SUM(h.qatnashchilar_soni),0) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                    AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s) AS qatnashchilar,
                   (SELECT COUNT(*) FROM murojaat m WHERE m.mahalla_id=mahalla.id AND m.sana BETWEEN %s AND %s) AS murojaat_soni
            FROM mahalla JOIN tuman ON mahalla.tuman_id=tuman.id
            WHERE tuman.viloyat_id=%s ORDER BY tuman.id, mahalla.mahalla_nomi
        """
        params = [start, end] * 3 + [viloyat_id]

    with connection.cursor() as cur:
        cur.execute(sql, params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    for r in rows:
        r['targibot_soni']  = int(r['targibot_soni'] or 0)
        r['qatnashchilar']  = int(r['qatnashchilar'] or 0)
        r['murojaat_soni']  = int(r['murojaat_soni'] or 0)
        # 1000 targ'ibot o'tkazilsa 1 murojaat bo'lishi ko'rsatkichi
        t = r['targibot_soni']
        m = r['murojaat_soni']
        r['nisbat'] = round(m / t, 2) if t > 0 else None  # murojaat/targ'ibot

    return Response({'rows': rows, 'start': start, 'end': end, 'role': role})


# ── 8. Xavfli mahallalar ─────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def xavfli_mahallalar(request):
    """
    Murojaat ko'p lekin targ'ibot kam bo'lgan mahallalar/viloyatlar.
    """
    start = request.GET.get('start', date(date.today().year, date.today().month, 1).isoformat())
    end   = request.GET.get('end',   date.today().isoformat())
    role  = request.user.role

    if role == 'respublika':
        sql = """
            SELECT v.nomi AS nomi, NULL AS tuman_nomi,
                   (SELECT COUNT(*) FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
                    JOIN tuman t ON m.tuman_id=t.id
                    WHERE t.viloyat_id=v.id AND h.status=2 AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s) AS targibot_soni,
                   (SELECT COUNT(*) FROM murojaat mr WHERE mr.viloyat_id=v.id AND mr.sana BETWEEN %s AND %s) AS murojaat_soni
            FROM viloyat v
            HAVING murojaat_soni > 0
            ORDER BY murojaat_soni DESC, targibot_soni ASC
        """
        params = [start, end] * 2
    else:
        viloyat_id = request.user.viloyat_id
        sql = """
            SELECT mahalla.mahalla_nomi AS nomi, tuman.tuman_nomi,
                   (SELECT COUNT(*) FROM hisobot h WHERE h.mahalla_id=mahalla.id AND h.status=2
                    AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s) AS targibot_soni,
                   (SELECT COUNT(*) FROM murojaat m WHERE m.mahalla_id=mahalla.id AND m.sana BETWEEN %s AND %s) AS murojaat_soni
            FROM mahalla JOIN tuman ON mahalla.tuman_id=tuman.id
            WHERE tuman.viloyat_id=%s
            HAVING murojaat_soni > 0
            ORDER BY murojaat_soni DESC, targibot_soni ASC
            LIMIT 100
        """
        params = [start, end] * 2 + [viloyat_id]

    with connection.cursor() as cur:
        cur.execute(sql, params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    for r in rows:
        r['targibot_soni'] = int(r['targibot_soni'] or 0)
        r['murojaat_soni'] = int(r['murojaat_soni'] or 0)
        # Xavf darajasi: murojaat bor, targ'ibot kam = yuqori xavf
        t = r['targibot_soni']
        m = r['murojaat_soni']
        if m > 0 and t == 0:
            r['daraja'] = 'yuqori'
        elif m > 0 and t < m * 5:
            r['daraja'] = 'o\'rta'
        else:
            r['daraja'] = 'past'

    return Response({'rows': rows, 'start': start, 'end': end, 'role': role})


# ── 9. Oylik dinamika ─────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def oylik_dinamika(request):
    """
    Yil bo'yicha oylik: targ'ibot soni, murojaat soni.
    Respublika: umumiy yoki viloyat filtri.
    """
    yil  = request.GET.get('yil', str(date.today().year))
    role = request.user.role

    if role == 'respublika':
        vid = request.GET.get('viloyat')
        if vid:
            targibot_sql = """
                SELECT MONTH(h.qushilgan_vaqt) AS oy, COUNT(*) AS soni,
                       COALESCE(SUM(h.qatnashchilar_soni),0) AS qatnashchilar
                FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
                JOIN tuman t ON m.tuman_id=t.id
                WHERE t.viloyat_id=%s AND h.status=2 AND YEAR(h.qushilgan_vaqt)=%s
                GROUP BY MONTH(h.qushilgan_vaqt)
            """
            murojaat_sql = """
                SELECT MONTH(sana) AS oy, COUNT(*) AS soni
                FROM murojaat WHERE viloyat_id=%s AND YEAR(sana)=%s
                GROUP BY MONTH(sana)
            """
            t_params = [int(vid), int(yil)]
            m_params = [int(vid), int(yil)]
        else:
            targibot_sql = """
                SELECT MONTH(qushilgan_vaqt) AS oy, COUNT(*) AS soni,
                       COALESCE(SUM(qatnashchilar_soni),0) AS qatnashchilar
                FROM hisobot WHERE status=2 AND YEAR(qushilgan_vaqt)=%s
                GROUP BY MONTH(qushilgan_vaqt)
            """
            murojaat_sql = """
                SELECT MONTH(sana) AS oy, COUNT(*) AS soni
                FROM murojaat WHERE YEAR(sana)=%s GROUP BY MONTH(sana)
            """
            t_params = [int(yil)]
            m_params = [int(yil)]
    else:
        vid = request.user.viloyat_id
        targibot_sql = """
            SELECT MONTH(h.qushilgan_vaqt) AS oy, COUNT(*) AS soni,
                   COALESCE(SUM(h.qatnashchilar_soni),0) AS qatnashchilar
            FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
            JOIN tuman t ON m.tuman_id=t.id
            WHERE t.viloyat_id=%s AND h.status=2 AND YEAR(h.qushilgan_vaqt)=%s
            GROUP BY MONTH(h.qushilgan_vaqt)
        """
        murojaat_sql = """
            SELECT MONTH(sana) AS oy, COUNT(*) AS soni
            FROM murojaat WHERE viloyat_id=%s AND YEAR(sana)=%s
            GROUP BY MONTH(sana)
        """
        t_params = [vid, int(yil)]
        m_params = [vid, int(yil)]

    OY_NOMI = ['','Yanvar','Fevral','Mart','Aprel','May','Iyun',
               'Iyul','Avgust','Sentabr','Oktabr','Noyabr','Dekabr']

    with connection.cursor() as cur:
        cur.execute(targibot_sql, t_params)
        targibot_map = {r[0]: {'soni': int(r[1]), 'qatnashchilar': int(r[2])} for r in cur.fetchall()}
        cur.execute(murojaat_sql, m_params)
        murojaat_map = {r[0]: int(r[1]) for r in cur.fetchall()}

    rows = []
    for oy in range(1, 13):
        rows.append({
            'oy': oy,
            'oy_nomi': OY_NOMI[oy],
            'targibot_soni': targibot_map.get(oy, {}).get('soni', 0),
            'qatnashchilar': targibot_map.get(oy, {}).get('qatnashchilar', 0),
            'murojaat_soni': murojaat_map.get(oy, 0),
        })

    return Response({'rows': rows, 'yil': yil})


# ── 10. Haftalik holat ────────────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def haftalik_holat(request):
    """
    Oxirgi 12 hafta bo'yicha: hafta boshidan targ'ibot va murojaat soni.
    """
    role = request.user.role

    if role == 'respublika':
        vid = request.GET.get('viloyat')
        if vid:
            targibot_sql = """
                SELECT YEARWEEK(qushilgan_vaqt, 1) AS hafta,
                       MIN(DATE(qushilgan_vaqt)) AS hafta_boshlash,
                       COUNT(*) AS soni,
                       COALESCE(SUM(qatnashchilar_soni),0) AS qatnashchilar
                FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
                JOIN tuman t ON m.tuman_id=t.id
                WHERE t.viloyat_id=%s AND h.status=2 AND qushilgan_vaqt >= DATE_SUB(CURDATE(), INTERVAL 12 WEEK)
                GROUP BY YEARWEEK(qushilgan_vaqt, 1) ORDER BY hafta DESC LIMIT 12
            """
            murojaat_sql = """
                SELECT YEARWEEK(sana, 1) AS hafta, COUNT(*) AS soni
                FROM murojaat WHERE viloyat_id=%s AND sana >= DATE_SUB(CURDATE(), INTERVAL 12 WEEK)
                GROUP BY YEARWEEK(sana, 1)
            """
            t_params = [int(vid)]
            m_params = [int(vid)]
        else:
            targibot_sql = """
                SELECT YEARWEEK(qushilgan_vaqt, 1) AS hafta,
                       MIN(DATE(qushilgan_vaqt)) AS hafta_boshlash,
                       COUNT(*) AS soni,
                       COALESCE(SUM(qatnashchilar_soni),0) AS qatnashchilar
                FROM hisobot
                WHERE status=2 AND qushilgan_vaqt >= DATE_SUB(CURDATE(), INTERVAL 12 WEEK)
                GROUP BY YEARWEEK(qushilgan_vaqt, 1) ORDER BY hafta DESC LIMIT 12
            """
            murojaat_sql = """
                SELECT YEARWEEK(sana, 1) AS hafta, COUNT(*) AS soni
                FROM murojaat WHERE sana >= DATE_SUB(CURDATE(), INTERVAL 12 WEEK)
                GROUP BY YEARWEEK(sana, 1)
            """
            t_params = []
            m_params = []
    else:
        vid = request.user.viloyat_id
        targibot_sql = """
            SELECT YEARWEEK(h.qushilgan_vaqt, 1) AS hafta,
                   MIN(DATE(h.qushilgan_vaqt)) AS hafta_boshlash,
                   COUNT(*) AS soni,
                   COALESCE(SUM(h.qatnashchilar_soni),0) AS qatnashchilar
            FROM hisobot h JOIN mahalla m ON h.mahalla_id=m.id
            JOIN tuman t ON m.tuman_id=t.id
            WHERE t.viloyat_id=%s AND h.status=2 AND h.qushilgan_vaqt >= DATE_SUB(CURDATE(), INTERVAL 12 WEEK)
            GROUP BY YEARWEEK(h.qushilgan_vaqt, 1) ORDER BY hafta DESC LIMIT 12
        """
        murojaat_sql = """
            SELECT YEARWEEK(sana, 1) AS hafta, COUNT(*) AS soni
            FROM murojaat WHERE viloyat_id=%s AND sana >= DATE_SUB(CURDATE(), INTERVAL 12 WEEK)
            GROUP BY YEARWEEK(sana, 1)
        """
        t_params = [vid]
        m_params = [vid]

    with connection.cursor() as cur:
        cur.execute(targibot_sql, t_params)
        targibot_rows = cur.fetchall()
        cur.execute(murojaat_sql, m_params)
        murojaat_map = {r[0]: int(r[1]) for r in cur.fetchall()}

    rows = []
    for r in reversed(targibot_rows):
        hafta_key   = r[0]
        hafta_start = str(r[1]) if r[1] else ''
        rows.append({
            'hafta': hafta_key,
            'hafta_boshlash': hafta_start,
            'targibot_soni': int(r[2]),
            'qatnashchilar': int(r[3]),
            'murojaat_soni': murojaat_map.get(hafta_key, 0),
        })

    return Response({'rows': rows})


# ── Hamkor tashkilot hisoboti ─────────────────────────────────────────────────
@api_view(['GET'])
@permission_classes([IsAuthenticated])
def hamkor_tashkilot_hisobot(request):
    role       = request.user.role
    viloyat_id = request.user.viloyat_id
    start      = request.GET.get('start', '')
    end        = request.GET.get('end', '')
    sel_viloyat = request.GET.get('viloyat', '')

    date_filter = ''
    params = []

    if start and end:
        date_filter = "AND DATE(h.qushilgan_vaqt) BETWEEN %s AND %s"
        params = [start, end]

    if role == 'respublika':
        v_filter = "AND ht.viloyat_id = %s" if sel_viloyat else ""
        if sel_viloyat:
            params = [sel_viloyat] + params

        sql = f"""
            SELECT
                v.nomi            AS viloyat_nomi,
                ht.id             AS tashkilot_id,
                ht.nomi           AS tashkilot_nomi,
                ht.turi           AS tashkilot_turi,
                COUNT(DISTINCT hx.id)                        AS xodim_soni,
                COUNT(h.id)                                  AS targibot_soni,
                COALESCE(SUM(h.qatnashchilar_soni), 0)       AS qatnashchilar
            FROM hamkor_tashkilot ht
            JOIN viloyat v ON v.id = ht.viloyat_id
            LEFT JOIN hamkor_xodim hx ON hx.tashkilot_id = ht.id AND hx.is_active = 1
            LEFT JOIN hisobot h ON h.hamkor_xodim_id IN (
                SELECT id FROM hamkor_xodim WHERE tashkilot_id = ht.id
            ) AND h.status = 2 {date_filter}
            WHERE 1=1 {v_filter}
            GROUP BY ht.id, ht.nomi, ht.turi, v.nomi
            ORDER BY targibot_soni DESC
        """
    else:
        sql = f"""
            SELECT
                ht.id             AS tashkilot_id,
                ht.nomi           AS tashkilot_nomi,
                ht.turi           AS tashkilot_turi,
                COUNT(DISTINCT hx.id)                        AS xodim_soni,
                COUNT(h.id)                                  AS targibot_soni,
                COALESCE(SUM(h.qatnashchilar_soni), 0)       AS qatnashchilar
            FROM hamkor_tashkilot ht
            LEFT JOIN hamkor_xodim hx ON hx.tashkilot_id = ht.id AND hx.is_active = 1
            LEFT JOIN hisobot h ON h.hamkor_xodim_id IN (
                SELECT id FROM hamkor_xodim WHERE tashkilot_id = ht.id
            ) AND h.status = 2 {date_filter}
            WHERE ht.viloyat_id = %s
            GROUP BY ht.id, ht.nomi, ht.turi
            ORDER BY targibot_soni DESC
        """
        params = params + [viloyat_id]

    with connection.cursor() as cur:
        cur.execute(sql, params)
        cols = [c[0] for c in cur.description]
        rows = [dict(zip(cols, r)) for r in cur.fetchall()]

    # Umumiy ko'rsatkichlar
    totals = {
        'tashkilot_soni': len(rows),
        'targibot_soni':  sum(r['targibot_soni'] for r in rows),
        'qatnashchilar':  sum(r['qatnashchilar'] for r in rows),
        'xodim_soni':     sum(r['xodim_soni'] for r in rows),
    }

    return Response({'rows': rows, 'totals': totals})
